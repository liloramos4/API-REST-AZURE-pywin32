import os
import sys
import subprocess



# Comprobar si la carpeta 'shell' ya existe
if not os.path.exists('shell'):
    # Si no existe, crear un entorno virtual llamado "shell"
    subprocess.run([sys.executable, "-m", "venv", "shell"])

# Definir la ubicación del ejecutable de Python en el entorno virtual
venv_python = os.path.join("shell", "Scripts", "python")
if sys.platform == "linux":
    venv_python = os.path.join("shell", "bin", "python")

# Actualizar pip en el entorno virtual
subprocess.run([venv_python, "-m", "pip", "install", "--upgrade", "pip"])

# Instalar las dependencias especificadas en el archivo requirements.txt
subprocess.run([venv_python, "-m", "pip", "install", "-r", "requirements.txt"])

# Crear el segundo script que se ejecutará dentro del entorno virtual
with open("second_script.py", "w", encoding='utf-8') as f:
    f.write("""
# -*- coding: utf-8 -*-

import json
import base64
import win32com.client
import win32com.client as win32
import re
import win32api
import urllib.parse  # Importado para analizar la URL
import requests
import os
from docxtpl import DocxTemplate
from docx import Document
from docx.text.paragraph import Paragraph
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.oxml import parse_xml
from docx.oxml.ns import nsmap
import docx
import time
from collections import OrderedDict
import shutil
from win32com.client import constants, gencache
from win32com.client import constants as wdConst
from win32com.client.gencache import EnsureDispatch
from urllib.parse import unquote
from bs4 import BeautifulSoup
from PIL import Image


# Pregunta al usuario si quiere descargar todas las páginas y subpáginas de la Wiki
respuesta = input("¿Deseas descargar todas las páginas de la Wiki Azure o solo una página? Por favor, responde con ‘sí’ para descargar todas las páginas, o ‘no’ para descargar solo una página? (si/no): ")

if respuesta.lower() == 'si':
    def todaslaspaginas():
        import json
        import base64
        import win32com.client
        from win32com.client import constants
        import win32com.client as win32
        import re
        import win32api
        import urllib.parse  # Importado para analizar la URL
        import requests
        import os
        from docxtpl import DocxTemplate
        from docx import Document
        from docx.text.paragraph import Paragraph
        from docx.opc.constants import RELATIONSHIP_TYPE as RT
        from docx.oxml import parse_xml
        from docx.oxml.ns import nsmap
        import docx
        import time
        from collections import OrderedDict
        import sys
        import pythoncom
        
        
        def limpiar_cache_com():
            print("Iniciando limpieza de caché COM...")
            try:
                # Asegurar que se pueda escribir en la caché
                win32com.client.gencache.is_readonly = False
                gen_py_path = win32com.__gen_path__
                print(f"Limpieza de caché COM en: {gen_py_path}")
                
                # Verificar si el directorio `gen_py` existe
                if os.path.exists(gen_py_path):
                    shutil.rmtree(gen_py_path)  # Eliminar el directorio
                    print("✓ Caché COM eliminada exitosamente")
                else:
                    print("No se encontró el directorio `gen_py`. No es necesario limpiarlo.")
                
                # Confirmar si se eliminó correctamente
                if not os.path.exists(gen_py_path):
                    print("✓ Confirmación: El directorio `gen_py` fue eliminado correctamente.")
                else:
                    print("✗ Error: No se pudo eliminar completamente el directorio `gen_py`.")
                    return False  # Salida anticipada si no se elimina
                
                pythoncom.CoInitialize()  # Inicializar COM para evitar problemas
                return True
            except Exception as e:
                print(f"Error al limpiar la caché COM: {e}")
                return False

        def inicializar_word():
            try:
                print("Intentando inicializar Word normalmente...")
                pythoncom.CoInitialize()  # Inicializar COM antes de `EnsureDispatch`
                word_app = win32com.client.gencache.EnsureDispatch('Word.Application')
                print("✓ Word inicializado correctamente")
                return word_app
            except AttributeError as e:
                # Detectar problemas de caché COM
                if "CLSIDToClassMap" in str(e) or "CLSIDToPackageMap" in str(e):
                    print("\\nProblema con la caché COM. Intentando solucionarlo...")
                    if limpiar_cache_com():
                        try:
                            print("\\nReintentando inicializar Word después de limpiar caché...")
                            # Intentamos regenerar la caché usando EnsureDispatch
                            word_app = win32com.client.gencache.EnsureDispatch('Word.Application')
                            print("✓ Word reinicializado exitosamente")
                            return word_app
                        except Exception as e2:
                            print(f"Error al reinicializar Word: {e2}")
                            raise
                else:
                    print(f"Error de atributo no relacionado con la caché: {e}")
                    raise
            except Exception as e:
                print(f"Error inesperado al inicializar Word: {e}")
                raise

        # Bloque principal para inicializar Word
        try:
            word_app = inicializar_word()
        except Exception as e:
            print(f"Fallo crítico: {e}. Intentando regenerar manualmente la caché.")
            try:
                pythoncom.CoInitialize()
                # Reforzamos la regeneración con EnsureDispatch
                win32com.client.gencache.EnsureDispatch('Word.Application')
                word_app = win32com.client.Dispatch('Word.Application')
            except Exception as e2:
                print(f"No se pudo generar la caché COM: {e2}")
                print("Considera reinstalar `pywin32` para resolver problemas de compatibilidad.")



        # Inicializa las variables globales
        stored_wiki_url = None
        stored_personal_access_token = None

        def sanitize_placeholder(placeholder):
            # Replace spaces and invalid characters with underscores
            sanitized = re.sub(r'[^\w]', '_', placeholder)
            # Remove leading digits and underscores to ensure a valid variable name
            sanitized = re.sub(r'^\d+|_', '', sanitized)
            return sanitized


        def update_toc(docx_file):
            try:
                word = win32com.client.DispatchEx("Word.Application")
                doc = word.Documents.Open(docx_file)
                doc.TablesOfContents(1).Update()
                doc.Close(SaveChanges=True)
                word.Quit()
            except Exception as e:
                print(f"An error occurred while updating the table of contents: {e}")


        def get_page_content(url):
            content_url = url + "?api-version=7.0&includeContent=true"
            response = requests.get(content_url, headers=headers)

            if response.status_code == 200:
                return json.loads(response.text)['content']
            else:
                print(f"Error al obtener el contenido de la página: {response.status_code}")
                return "Contenido no disponible"

        def extract_pages_recursive(page, level=1):
            if not page:
                return []

            original_content = get_page_content(page['url'])
            content = get_page_content(page['url'])                    
            # convertilos dobles >> << en parentesis también
            content = re.sub(r'(<<|>>|<|>)', lambda m: {'<<':'((','>>':'))','<':'(','>':')'}[m.group(0)], content)
            content = re.sub(r'<b><span style="color:([^>]*)">([^<]*)</span></b>', r'(b)(span style="color:\\1")\\2(/span)(/b)', content, flags=re.IGNORECASE)
            content = re.sub(r'\*\*\s*\(/span\)\s*\|', '** |', content) 
            content = re.sub(r'<center>(.*?)</center>', r'\\1', content)
            content = re.sub(r'<code>```[^\\n]*\\n', '```\\n', content)
            content = re.sub(r'<br>(.*?)</br>', r'\\1', content)
            content = re.sub(r'(<span style="color:([^>]*)">([^<]*?))\\n', r'\\1(/span)\\n', content)
            content = re.sub(r'<br>', '\\n', content)
            content = re.sub(r'TO_DO: @<([A-Fa-f0-9-]+)>', r'TO_DO: \\1', content)
            content = re.sub(r'<Lista> @<([^>]+)>', r'Lista @\\1', content)
            content = re.sub(r'\.(PNG|JPG|JPEG|GIF)', lambda x: x.group().lower(), content)
            content = re.sub(r'\|\s*(\(span style="color:[^)]+\))\s*([^|]+?)\s*(?=\|)', r'| \\1 \\2(/span)', content, flags=re.IGNORECASE)
            content = re.sub(r'(\(span style="color:[^)]+\)[^|]*?)\s*\|', r'\\1 (/span)|', content)
            content = re.sub(r'\|\(/span\)', '|', content)
            content = re.sub(r'(\(/span\))(\s*\(/span\)\s*\|)', r'\\1|', content) 
            content = re.sub(r'\(/b\)\s*\(/span\)\|', '(/b)|', content)
            content = re.sub(r'\(/b\)\s*\(/span\)\s*\|', '(/b) |', content)
            # (/span) eliminado  ya que hay espacio development/DevOps tools (/span)| ejemplo1 |
            content = re.sub(r'\s+\(/span\)\|', '|', content)
            expresion_regular = r'<PRE[^>]*><CODE[^>]*><DIV><B>(GET [^<]+)</B><BR/></DIV></CODE></PRE>'
            content = re.sub(expresion_regular, r'negrit````\\n\\1\\n````negrit\\n\\n', content)

            # Expresión regular para extraer el texto dentro de la etiqueta DIV (sin cambios)
            content = re.sub(r'<DIV style="[^"]*">(.*?)</DIV>\s*', r'\\1\\n', content)

            # Convertir etiquetas HTML a paréntesis
            parts = re.split(r'(```.*?```)', content, flags=re.DOTALL)
            content = ''.join(
                re.sub(r'<([^>]*)>', r'(\\1)', part) if not part.startswith('```') else part
                for part in parts
            )


            # Dentro de bloques de código especifica signos < y > 
            content = re.sub(
                r'```.*?```',
                lambda m: m.group(0).replace('<', '&lt;').replace('>', '&gt;'),
                content,
                flags=re.DOTALL
            )

            
            # Formatear encabezados de tabla y eliminar saltos de línea adicionales
            content = re.sub(r'\(TR\).*?\(TH.*?\)(.*?)\(/TH\)\s*\\n*\s*\(TH.*?\)(.*?)\(/TH\).*?\(/TR\).*?\(/THEAD\)', r'| \\1 | \\2 |\\n|--|--|\\n', content)

            # Formatear celdas de tabla
            content = re.sub(r'\(TR\).*?\(TD.*?\)(.*?)\(/TD\)\s*\(TD.*?\)(.*?)\(/TD\).*?\(/TR\)', r'| \\1 | \\2 |\\n', content)

            # Eliminar elementos adicionales
            content = re.sub(r'\(DIV\)|\(TABLE\)|\(THEAD\)|\(TBODY\)|\(TR\)|\(BR/\)', '', content)

            # elimina (DIV )(TABLE )
            content = re.sub(r'\(DIV style="color: rgb\(0, 0, 0\);font-family: &amp;quot;font-size: 14px;font-style: normal;font-weight: 400;letter-spacing: normal;text-align: start;text-indent: 0px;text-transform: none;white-space: normal;word-spacing: 0px"\)|\(TABLE style="border-collapse: collapse;border-spacing: 0px;margin: 1rem 0px 0px;padding: 0px;border: 0px;width: 860px;table-layout: auto;font-size: 0.875rem;color: rgb\(23, 23, 23\);font-family: &amp;quot"\)', '', content)

            content = re.sub(r'\(/TBODY\)|\(/TABLE\)|\(/DIV\)', '', content)

            content = re.sub(r'\(CODE style=".*?"\)\(SPAN style=".*?"\)(.*?)\(/SPAN\)\(/CODE\)', r'``\\1``', content)

            content = re.sub(r'\(CODE style=".*?"\)|\(/CODE\)', '`', content)

            content = re.sub(r'\((/?SPAN)\)', '', content, flags=re.IGNORECASE)

            # Anidar y aplicar ambas expresiones regulares en una sola línea
            content = re.sub(r'\|>', '|', re.sub(r'(\|[^|]*)(?=>)', r'\\1|', content))

            # Añade saltos de linea después de un pipe si se ve bloque de código
            content = re.sub(r'(\|\s*)\\n(```)', r'\\1\\n\\n\\2', content)

             
            # expresión regular para corregir los encabezados de la tabla Markdown       
            content = re.sub(r'(\|\s*\w+\s*\|\s*\w+\s*\|)\s*(\|\-\-\|\-\-\|)', r'\\1\\n\\2', content)

            # Expresión regular para detectar las filas de la tabla y ajustar los separadores
            content = re.sub(r'(\|\s*[^|]+\s*\|\s*[^|]+\s*\|\s*[^|]+\s*\|\s*[^|]+\s*\|)\s*\|--\|--\|--\|--\|--\|', r'\\1\\n|--|--|--|--|', content)

            # Modificamos la expresión regular fix tabla y header encabezado
            content = re.sub(r'\|\s*([^\\n|]+)\s*\\n\s*\|\s*([^\\n|]+)\s*\|\s*\|\s*---\s*---\s*\|', r'| \\1 | \\2 |\\n| --- | --- |', content)

            # Eliminar porcentaje si va acompañado de guión bajo
            content = re.sub(r'%_', '_', content)

            # Ajustes de etiquetas HTML espacios innecesarios
            content = re.sub(r'\s*((</?(?:b|span[^>]*)>))\s*', r'\\1', content)

            # Expresión regular para modificar la tabla en una sola línea
            content = re.sub(r'(\| Purpose User  \|)|(\|--\|--\|--\|)', lambda m: '| Purpose | User  |' if m.group(1) else '|--|--|--|--|', content)

            # (b)(span style="color:red")IMPORTANT:(/b)| realiza un salto de párrafo ya que no tiene tuberia al principio
            content = re.sub(r'^([^|]+)(\|)', r'\\1\\n\\2', content, flags=re.MULTILINE)

            # Paso 2 actualizado: Asegurarse de que las celdas vacías antes del encabezado tienen un salto de línea
            content = re.sub(r'(\|[ \t]*\|[ \t]*\|[ \t]*\|)(--\|--\|--\|--\|)', r'\\1\\n|\\2', content)

            # paso3 identifica https y http luego quita acentos graves innecesarios  
            content = re.sub(r'`(https?://[^\s`]+)`', r'\\1', content)
            
            # Cuando la etiqueta (b)span style está pega a dos puntos
            content = re.sub(r':(?=\(b\)|\(span style="color:[^"]+"\))', ':\\n', content)
            
            # Expresión regular que añade `(/span)` si falta después de `(/b)`, con cualquier color en el estilo
            content = re.sub(r"(\(b\)\(span style=\\"color:[a-zA-Z]+\\"\).*?\(\/b\))(?!\(\/span\))", r"\\1(/span)", content)
            
            # Para espacios dobles después de la etiqueta de apertura
            content = re.sub(r'(\(span [^)]+\).*?)(\s{2,})', r'\\1(/span)\\2', content)  # Para espacios dobles
  
            # Para finales de línea sin (/span) agrega 1 (/span)Para finales de línea sin (/span)
            content = re.sub(r'\((span [^)]+)\)([^(/\\n]+?)(?<!(/span))(?=\s*\||(/b)|(\\n)|$)', r'(\\1)\\2 (/span)', content) 
  
            # Expresión regular para buscar `(/b)(/span)(/span)` y eliminar solo el último `(/span)`
            content = re.sub(r"\(/b\)\(/span\)\(/span\)", r"(/b)(/span)", content)
            
            # Expresión regular para quitar '(/span)' que está cerca de una tubería '|'
            content = re.sub(r'\|(\s*)\(/span\)', r'|', content)
            
            # Usamos re.sub para agregar un salto de línea después de los dos puntos y antes de "negrit"
            content = re.sub(r":(negrit````)", r":\\n\\1", content)
            
            # Usamos re.sub para agregar un salto de línea si detrás de los dos puntos hay palabras y luego tres guiones
            content = re.sub(r"(\w+:\s*)---", r"\\1\\n---", content)
            
            # Expresión regular para quitar '(/span)' que está cerca de un '.com'
            content = re.sub(r'(\.com)\s*\(/span\)', r'\\1', content)
            
            # Usar re.sub para eliminar los saltos de línea intermedios SOLO cuando hay 3 guiones
            content = re.sub(r'\\n{2,}(?=---)', '\\n', content)
            
            # Insertar saltos de párrafos después de los dos puntos antes de la imagen HTML
            content = re.sub(r':(?!\\n)(?=\(a href="[^"]+"\)\(img src="[^"]+"[^>]*\/\)\(\/a\))', ':\\n', content)
            
            # Genera 1 espacio de más después de los 2 puntos
            content = re.sub(r'(?<!`\*\*):(?!\s|\*)(?![a-zA-Z0-9]*\*\*`)(?=[A-Za-z0-9])', ': ', content)

            # inserta un salto de linea en el caso que no haya ningún espacio entre los dos puntos y el texto 
            content = re.sub(r'(:)(?=(?!\*\*)(?:[\w`]|```))', r'\\1\\n', content)

            # Funciona todo alinea a la misma altura un asterisco con otro asterisco
            content = re.sub(r'(^.*:\s*\\n)(\*)(\s*\S.*)', r'\\1    \\2\\3', content, flags=re.MULTILINE)
            
            # Esta expresión regular quitar inserción de saltos de linea innecesarios dentro bloques de código márkdowns
            content = re.sub(r'```(.*?)\\n(.*?)\\n([ \\t]*\\|.*?)```', r'```\\1\\n\\2 \\3```', content, flags=re.DOTALL)

            # expresión regular que autoajusta espacios necesarios cabeceras de las tablas markdown
            # Función de reemplazo
            def replacer(m):
                # Si coincidió la parte separadora (?P<sep>)
                if m.group("sep") is not None:
                    # Convertir la línea separadora al formato "|--|--|..." según cantidad de columnas
                    sep_line = m.group("sep").strip()
                    columns = [col for col in sep_line.split("|") if col.strip()]
                    return "|" + "|".join(["--"] * len(columns)) + "|"
                else:
                    # Sino es la parte de encabezado (?P<header>), quitamos espacios iniciales
                    return m.group("header").lstrip()

            # Expresión regular aplicada directamente con re.sub()
            content = re.sub(r"^(?P<sep>\s*\|(?:\s*[-:]+\s*\|)+\s*)$|^(?P<header>\s+\|(?:[^|]*\|)+.*)$", 
                                        replacer, 
                                        content, 
                                        flags=re.MULTILINE)
 
            
            # Ajusta la r la pone de bajo en el código YAML 
            content = re.sub(r'(```[\s\S]*?```|\(code\)[\s\S]*?\(/code\))', lambda m: re.sub(r'(^[ \t]*-[ \t]*script:[ \t]*\\|)[ \t]*(#.*$)', lambda a: f"{a.group(1)}\\n{' ' * a.group(1).find('r')}{a.group(2).strip()}", m.group(0), flags=re.MULTILINE), content, flags=re.DOTALL)

            # 1) Aplana pipes dentro de (code)…(/code) sin tocar los comentarios
            content = re.sub(r'(```[\s\S]*?```|\(code\)[\s\S]*?\(/code\))', lambda m: m.group(0).replace("\\n|", " |").replace("| \\n", " | "), content, flags=re.DOTALL)

            # para identar yaml correctamente y calcula los espacios
            content = re.sub(
                r'(```[\s\S]*?```|\(code\)[\s\S]*?\(/code\))', 
                lambda m: re.sub(
                    r'(-)[ \t]+(script)[ \t]*(:)[ \t]*(?=\\|)',
                    r'\\1 \\2\\3 ',
                    m.group(0)
                ), 
                content, 
                flags=re.DOTALL
            )
            # Mantiene todo en orden YAML dentro de los bloques de código
            content = re.sub(r'(```[\s\S]*?```|\(code\)[\s\S]*?\(/code\))', lambda m: re.sub(r'(^[ \t]*-[ \t]*script:[ \t]*\|)[ \t]*(#.*$)', lambda a: f"{a.group(1)}\\n{' ' * a.group(1).find('r')}{a.group(2).strip()}", m.group(0), flags=re.MULTILINE), content, flags=re.DOTALL)


            # Solamente calcula y  convierte de un signo > a parentesis blockquote. Convierte todos los blockquote borde izquierdo en parentesis para que se pueda procesar bien
            lines = content.strip('\\n').splitlines()
            result = []
            prev_blockquote = False

            for line in lines:
                if re.match(r'^\s*>+\s*(.*)', line):
                    result.append(re.sub(r'^\s*>+\s*(.*)', r') \\1', line))
                    prev_blockquote = True
                elif re.match(r'^\s*\|\s*.*', line) and prev_blockquote:
                    result.append(re.sub(r'^\s*\|\s*(.*)', r') \\1', line))
                    prev_blockquote = False  # solo permitir uno seguido
                else:
                    result.append(line)
                    prev_blockquote = False  # se corta la cadena

            content = '\\n'.join(result) 
                                     
            info = {
                'name': page['path'],
                'short_name': page['path'].split('/')[-1],
                'url': page['url'],
                'original_content': original_content,
                'content': content ,
                'level': level,
                'subpages': []
            }


            page_info = [info]

            if 'subPages' in page:
                for sub_page in page['subPages']:
                    page_info.extend(extract_pages_recursive(sub_page, level+1))

            return page_info

        def extract_url_values(url):
            regex = r"https://dev.azure.com/(?P<organization>[^/]+)/(?P<project>[^/]+)/_wiki/wikis/(?P<wiki>[^/]+)/.*"
            match = re.search(regex, url)
            if match:
                return match.groupdict()
            else:
                return None


        def extract_placeholders(template_path):
            doc = Document(template_path)
            placeholders = []

            for element in doc.element.body.iter():
                if element.tag.endswith('p'):
                    paragraph = Paragraph(element, doc)
                    if hasattr(paragraph, 'runs'):
                        # Concatenate the text of adjacent runs
                        full_text = ''.join([run.text for run in paragraph.runs])
                        
                        # Find placeholders in the full text
                        start_index = full_text.find('{{')
                        end_index = full_text.find('}}')
                        while start_index != -1 and end_index != -1:
                            placeholder = full_text[start_index+2:end_index].strip()
                            placeholders.append(placeholder)
                            
                            # Find the next placeholder
                            start_index = full_text.find('{{', end_index)
                            end_index = full_text.find('}}', end_index+2)

                elif element.tag.endswith('tc'):  # Check for table cell (td)
                    for p in element.iterchildren('{%s}p' % nsmap['w']):
                        paragraph = Paragraph(p, doc)
                        if hasattr(paragraph, 'runs'):
                            # Concatenate the text of adjacent runs
                            full_text = ''.join([run.text for run in paragraph.runs])
                            
                            # Find placeholders in the full text
                            start_index = full_text.find('{{')
                            end_index = full_text.find('}}')
                            while start_index != -1 and end_index != -1:
                                placeholder = full_text[start_index+2:end_index].strip()
                                placeholders.append(placeholder)
                                
                                # Find the next placeholder
                                start_index = full_text.find('{{', end_index)
                                end_index = full_text.find('}}', end_index+2)

            # Remove duplicates from the list of placeholders while maintaining the order of the elements
            placeholders = list(OrderedDict.fromkeys(placeholders))

            return placeholders


        def create_context(page_info, placeholders):
            context = {}
            title_index = 1
            for placeholder in placeholders:
                if title_index < len(page_info):
                    page = page_info[title_index]
                    # Agregar valores al contexto tanto para marcadores de posición de título como para marcadores de posición de contenido
                    if placeholder.endswith('_content'):
                        if page['content'] != 'No hay contenido':
                            context[placeholder] = page['content']
                        title_index += 1
                    else:
                        context[placeholder] = page['short_name']
            return context

        # Pedir la información al usuario una sola vez y almacenarla en variables globales
        # Verificar si las credenciales ya se han almacenado
        if stored_wiki_url is None or stored_personal_access_token is None:
            wiki_url = input("Introduce la URL principal del portal Wiki de Azure: ")
            personal_access_token = input("Introduce tu token de acceso personal: ")
            # Aquí iría la lógica para verificar si las credenciales son válidas.
            # Si son válidas, las almacenamos en las variables globales.
            stored_wiki_url = wiki_url
            stored_personal_access_token = personal_access_token
        else:
            # Usar las credenciales almacenadas
            wiki_url = stored_wiki_url
            personal_access_token = stored_personal_access_token


        credentials = f":{personal_access_token}"
        encoded_credentials = base64.b64encode(credentials.encode('utf-8')).decode('utf-8')

        headers = {
            'Authorization': f'Basic {encoded_credentials}',
            "Content-Type": "application/json",
            "Accept": "application/json"
        }

        # Decodifica la URL para manejar caracteres especiales
        decoded_wiki_url = urllib.parse.unquote(wiki_url)
        # Usa regex para extraer la organización, el proyecto y el wiki de la URL decodificada
        url_match = re.search(r"https://dev.azure.com/(?P<organization>[^/]+)/(?P<project>[^/]+)/_wiki/wikis/(?P<wiki>[^/]+)/.*", decoded_wiki_url)
        if url_match:
            organization = url_match.group('organization')
            project = url_match.group('project')
            wiki = url_match.group('wiki')
        else:
            print("The URL provided is not valid.")
            exit()

        url = f"https://dev.azure.com/{organization}/{project}/_apis/wiki/wikis/{wiki}/pages?api-version=7.0&recursionLevel=full"

        response = requests.get(url, headers=headers)

        if response.status_code == 200:
            print("Successful response http 200.")
            root_page = json.loads(response.text)
        else:
            print(f"Error al obtener la página raíz de la wiki: {response.status_code}")
            root_page = {}

        page_info = extract_pages_recursive(root_page)


        # Save the .md file with original content
        md_filename_original = 'htmlymd.md'
        with open(md_filename_original, 'w', encoding='utf-8') as f:
            for page in page_info:
                # Get the title and original content of the page
                title = page['name'].split('/')[-1]
                original_content = page['original_content'].strip()

                # Only write the title and original content to the Markdown file if the original content is not empty
                if original_content and not (title.startswith("#") and "No hay contenido" in title) and page['name'] != "/":
                    f.write(f'# {title}\\n')
                    f.write(original_content)
                    f.write('\\n\\n')


        # Save the .md file
        md_filename = 'todosmd.md'
        with open('todosmd.md', 'w', encoding='utf-8') as f:
            for page in page_info:
                # Get the title and content of the page
                title = page['name'].split('/')[-1]
                content = page['content'].strip()

                # Only write the title and content to the Markdown file if the content is not empty
                if content and not (title.startswith("#") and "No hay contenido" in title) and page['name'] != "/":
                    f.write(f'# {title}\\n')
                    f.write(content)
                    f.write('\\n\\n')


        # Read the Markdown content
        with open(md_filename, 'r', encoding='utf-8') as f:
            markdown_content = f.read()

        # Extract the level 1 headings with regex
        title_pattern = r'^# (.*)'
        titles = re.findall(title_pattern, markdown_content, flags=re.M)


        # Generate the placeholders
        placeholders = [{'text': f'Titulo{i+1}', 'level': page['level']} for i, page in enumerate(page_info)]

        def add_titles_to_template(template_path, titles_info):
            doc = Document(template_path)

            # Find the element after which to insert the placeholders
            insert_element = None
            for element in doc.element.body.iter():
                if element.tag.endswith('p'):
                    paragraph = Paragraph(element, doc)
                    if hasattr(paragraph, 'runs'):
                        # Concatenate the text of adjacent runs
                        full_text = ''.join([run.text for run in paragraph.runs])
                        if 'REFERENCES' in full_text:
                            insert_element = paragraph
                            break

            if insert_element is not None:
                # Insert the titles before the element with the desired heading style
                for title_info in titles_info:
                    # Skip titles with name '/' or empty short_name
                    if 'name' in title_info and (title_info['name'] == '/' or not title_info['short_name']):
                        continue
                    sanitized_title = sanitize_placeholder(title_info["title"])
                    p = insert_element.insert_paragraph_before(f'{{{{{sanitized_title}}}}}')
                    # Calculate the maximum level in the JSON
                    max_level = max([title_info['level'] for title_info in titles_info])
                    # Generate the level_map dictionary dynamically
                    level_map = {i: f'Heading {i}' for i in range(1, max_level+1)}
                    heading_style = level_map.get(title_info["level"] - 1, 'Normal')
                    p.style = heading_style
                    # Insert a placeholder for the content below the title
                    p = insert_element.insert_paragraph_before(f'{{{{{sanitized_title}_content}}}}')

                # Insert a page break after the last title
                p = insert_element.insert_paragraph_before()
                run = p.add_run()
                run.add_break(docx.enum.text.WD_BREAK.PAGE)

            # Save the modified template
            doc.save('new_template.docx')


        # Generate the title information
        titles_info = [{'title': page['name'].split('/')[-1], 'level': page['level']} for page in page_info if page['content'].strip() and page['name'].split('/')[-1].strip()]


        # Call the function to create a new template and add the titles to it
        add_titles_to_template('plantilla.docx', titles_info)

        # Load the modified template
        doc = DocxTemplate("new_template.docx")

        # Extract the placeholders from the template
        placeholders = extract_placeholders("new_template.docx")


        # Create a new context
        context = create_context(page_info, placeholders)


        # Render template with dynamic context
        doc.render(context)

        # Save the generated document
        doc.save("documento_generado.docx")


        print("Wait one moment to finish work")
        # Wait for a few seconds to make sure the file is saved
        time.sleep(4)
        print(" Already ")

        # Check if the file was saved successfully
        if os.path.exists("documento_generado.docx"):
            print("The generated_document.docx file has been saved successfully.")
        else:
            print("No se pudo guardar el archivo documento_generado.docx.")

        # Get the directory of the script
        dir_path = os.path.dirname(os.path.realpath(__file__))

        # Build the file path
        docx_file = os.path.join(dir_path, "documento_generado.docx")

        # Call the function to update the table of contents in the generated document
        update_toc(docx_file)  


        def clean_attachments_folder():
            ###Elimina todos los archivos en la carpeta .attachments
            attachments_dir = '.attachments'
            if os.path.exists(attachments_dir):
                for file_name in os.listdir(attachments_dir):
                    file_path = os.path.join(attachments_dir, file_name)
                    try:
                        if os.path.isfile(file_path) or os.path.islink(file_path):
                            os.unlink(file_path)
                            print(f"Deleted file: {file_name}")
                    except Exception as e:
                        print(f"Failed to delete {file_name}: {e}")
            else:
                os.makedirs(attachments_dir)  # Si no existe, la creamos

        # Limpiar la carpeta .attachments antes de descargar nuevos archivos
        clean_attachments_folder()

        print("Preparandose para descargar azure wiki GIT")       
        
        list_url = f"https://dev.azure.com/{organization}/{project}/_apis/git/repositories/{wiki}/items?scopePath=/.attachments&recursionLevel=full&api-version=5.0"

        # Configura la autenticación
        headers = {
            'Authorization': f'Basic {base64.b64encode((":{}".format(personal_access_token)).encode()).decode()}'
        }

        # Crea la carpeta .attachments si no existe
        if not os.path.exists('.attachments'):
            os.makedirs('.attachments')

        # Realiza la solicitud GET para obtener la lista de archivos
        list_response = requests.get(list_url, headers=headers)

        # Procesa la respuesta
        if list_response.status_code == 200:
            response_json = list_response.json()
            if 'value' in response_json:
                files = response_json['value']
                folders = [file for file in files if file.get('isFolder')]
                for folder in folders:
                    print(f"Your folder is {folder['path']}")
                for file in files:
                    # Salta el elemento que representa la carpeta en sí
                    if file.get('isFolder'):
                        continue

                    # Descarga cada archivo
                    download_url = f"https://dev.azure.com/{organization}/{project}/_apis/git/repositories/{wiki}/items?path={file['path']}&api-version=5.0"
                    download_response = requests.get(download_url, headers=headers)
                    if download_response.status_code == 200:
                        file_name = os.path.basename(file['path'])
                        with open(f'.attachments/{file_name}', 'wb') as f:
                            f.write(download_response.content)
                        print(f"Successfully downloaded the file {file_name} de la carpeta {folder['path']}")
                    else:
                        print(f"Error al descargar {file['path']} de la carpeta {folder['path']}: {download_response.status_code}")
            else:
                print("La clave 'value' no está presente en la respuesta de la API.")
        else:
            print(f"Error al listar archivos: {list_response.status_code}. Ese código 404 no te preocupues simplemente en tu wiki no tienes ficheros imágenes de momento")

        pass
        
    todaslaspaginas()  # Llama a la función todaslaspaginas

else:

    def paginaconcreta2():
        import requests
        import json
        import base64
        import re
        from docxtpl import DocxTemplate
        import os
        from docx import Document
        from docx.text.paragraph import Paragraph
        from docx.opc.constants import RELATIONSHIP_TYPE as RT
        from docx.oxml import parse_xml
        from docx.oxml.ns import nsmap
        from collections import OrderedDict
        import docx
        import win32com.client
        import time
        import urllib.parse
        import subprocess
        import sys
        sys.coinit_flags = 0  # Asegura la inicialización correcta de COM
        import shutil
        import pythoncom
        from urllib.parse import unquote

        
        def limpiar_cache_com():
            print("Iniciando limpieza de caché COM...")
            try:
                # Asegurar que se pueda escribir en la caché
                win32com.client.gencache.is_readonly = False
                gen_py_path = win32com.__gen_path__
                print(f"Limpieza de caché COM en: {gen_py_path}")
                
                # Verificar si el directorio `gen_py` existe
                if os.path.exists(gen_py_path):
                    shutil.rmtree(gen_py_path)  # Eliminar el directorio
                    print("✓ Caché COM eliminada exitosamente")
                else:
                    print("No se encontró el directorio `gen_py`. No es necesario limpiarlo.")
                
                # Confirmar si se eliminó correctamente
                if not os.path.exists(gen_py_path):
                    print("✓ Confirmación: El directorio `gen_py` fue eliminado correctamente.")
                else:
                    print("✗ Error: No se pudo eliminar completamente el directorio `gen_py`.")
                    return False  # Salida anticipada si no se elimina
                
                pythoncom.CoInitialize()  # Inicializar COM para evitar problemas
                return True
            except Exception as e:
                print(f"Error al limpiar la caché COM: {e}")
                return False

        def inicializar_word():
            try:
                print("Intentando inicializar Word normalmente...")
                pythoncom.CoInitialize()  # Inicializar COM antes de `EnsureDispatch`
                word_app = win32com.client.gencache.EnsureDispatch('Word.Application')
                print("✓ Word inicializado correctamente")
                return word_app
            except AttributeError as e:
                # Detectar problemas de caché COM
                if "CLSIDToClassMap" in str(e) or "CLSIDToPackageMap" in str(e):
                    print("\\nProblema con la caché COM. Intentando solucionarlo...")
                    if limpiar_cache_com():
                        try:
                            print("\\nReintentando inicializar Word después de limpiar caché...")
                            # Intentamos regenerar la caché usando EnsureDispatch
                            word_app = win32com.client.gencache.EnsureDispatch('Word.Application')
                            print("✓ Word reinicializado exitosamente")
                            return word_app
                        except Exception as e2:
                            print(f"Error al reinicializar Word: {e2}")
                            raise
                else:
                    print(f"Error de atributo no relacionado con la caché: {e}")
                    raise
            except Exception as e:
                print(f"Error inesperado al inicializar Word: {e}")
                raise

        # Bloque principal para inicializar Word
        try:
            word_app = inicializar_word()
        except Exception as e:
            print(f"Fallo crítico: {e}. Intentando regenerar manualmente la caché.")
            try:
                pythoncom.CoInitialize()
                # Reforzamos la regeneración con EnsureDispatch
                win32com.client.gencache.EnsureDispatch('Word.Application')
                word_app = win32com.client.Dispatch('Word.Application')
            except Exception as e2:
                print(f"No se pudo generar la caché COM: {e2}")
                print("Considera reinstalar `pywin32` para resolver problemas de compatibilidad.")

       
        # Inicializa las variables globales
        stored_wiki_url = None
        stored_personal_access_token = None

        def sanitize_placeholder(placeholder):
            # Replace spaces and invalid characters with underscores
            sanitized = re.sub(r'[^\w]', '_', placeholder)
            # Remove leading digits and underscores to ensure a valid variable name
            sanitized = re.sub(r'^\d+|_', '', sanitized)
            return sanitized

        def update_toc(docx_file):
            # No eliminamos gen_py aquí, ya lo hemos hecho manualmente
            # Continuar con las tareas de Word
            try:
                print("Iniciando la aplicación de Word y actualizando la tabla de contenido...")
                word = win32com.client.gencache.EnsureDispatch('Word.Application')
                word.Visible = False
                doc = word.Documents.Open(docx_file)
                doc.TablesOfContents(1).Update()
                doc.Close(SaveChanges=True)
                word.Quit()
                print("Documento de Word actualizado correctamente.")
            except Exception as e:
                print(f"Error al intentar iniciar o procesar la aplicación de Word: {e}")

        def get_page_content(url, headers):
            content_url = url + "?api-version=7.0&includeContent=true"
            response = requests.get(content_url, headers=headers)

            if response.status_code == 200:
                # Convierte la respuesta en un objeto JSON
                response_json = json.loads(response.text)
                # Devuelve el valor del campo 'content'
                return response_json['content']
            else:
                print(f"Error al obtener el contenido de la página: {response.status_code}")
                return "Contenido no disponible"

        def extract_pages_recursive(page, headers=None, level=1):
            if not page:
                return []

            original_content = get_page_content(page['url'], headers)
            content = get_page_content(page['url'], headers)
            # Replace <span> tags with plain text
            content = re.sub(r'<b><span style="color:([^>]*)">([^<]*)</span></b>', r'(b)(span style="color:\\1")\\2(/span)(/b)', content, flags=re.IGNORECASE)
            content = re.sub(r'\*\*\s*\(/span\)\s*\|', '** |', content) 
            content = re.sub(r'<center>(.*?)</center>', r'\\1', content)
            content = re.sub(r'<code>```[^\\n]*\\n', '```\\n', content)
            content = re.sub(r'<br>(.*?)</br>', r'\\1', content)
            content = re.sub(r'(<span style="color:([^>]*)">([^<]*?))\\n', r'\\1(/span)\\n', content)
            content = re.sub(r'<br>', '\\n', content)
            content = re.sub(r'TO_DO: @<([A-Fa-f0-9-]+)>', r'TO_DO: \\1', content)
            content = re.sub(r'<Lista> @<([^>]+)>', r'Lista @\\1', content)
            content = re.sub(r'\.(PNG|JPG|JPEG|GIF)', lambda x: x.group().lower(), content)
            content = re.sub(r'\|\s*(\(span style="color:[^)]+\))\s*([^|]+?)\s*(?=\|)', r'| \\1 \\2(/span)', content, flags=re.IGNORECASE)
            content = re.sub(r'(\(span style="color:[^)]+\)[^|]*?)\s*\|', r'\\1 (/span)|', content)
            content = re.sub(r'\|\(/span\)', '|', content)
            content = re.sub(r'(\(/span\))(\s*\(/span\)\s*\|)', r'\\1|', content) 
            content = re.sub(r'\(/b\)\s*\(/span\)\|', '(/b)|', content)
            content = re.sub(r'\(/b\)\s*\(/span\)\s*\|', '(/b) |', content)
            # (/span) eliminado  ya que hay espacio development/DevOps tools (/span)| ejemplo1 |
            content = re.sub(r'\s+\(/span\)\|', '|', content)
            expresion_regular = r'<PRE[^>]*><CODE[^>]*><DIV><B>(GET [^<]+)</B><BR/></DIV></CODE></PRE>'
            content = re.sub(expresion_regular, r'negrit````\\n\\1\\n````negrit\\n\\n', content)

            # Expresión regular para extraer el texto dentro de la etiqueta DIV (sin cambios)
            content = re.sub(r'<DIV style="[^"]*">(.*?)</DIV>\s*', r'\\1\\n', content)

            # Convertir etiquetas HTML a paréntesis
            parts = re.split(r'(```.*?```)', content, flags=re.DOTALL)
            content = ''.join(
                re.sub(r'<([^>]*)>', r'(\\1)', part) if not part.startswith('```') else part
                for part in parts
            )


            # Dentro de bloques de código especifica signos < y > 
            content = re.sub(
                r'```.*?```',
                lambda m: m.group(0).replace('<', '&lt;').replace('>', '&gt;'),
                content,
                flags=re.DOTALL
            )

            # Formatear encabezados de tabla y eliminar saltos de línea adicionales
            content = re.sub(r'\(TR\).*?\(TH.*?\)(.*?)\(/TH\)\s*\\n*\s*\(TH.*?\)(.*?)\(/TH\).*?\(/TR\).*?\(/THEAD\)', r'| \\1 | \\2 |\\n|--|--|\\n', content)

            # Formatear celdas de tabla
            content = re.sub(r'\(TR\).*?\(TD.*?\)(.*?)\(/TD\)\s*\(TD.*?\)(.*?)\(/TD\).*?\(/TR\)', r'| \\1 | \\2 |\\n', content)

            # Eliminar elementos adicionales
            content = re.sub(r'\(DIV\)|\(TABLE\)|\(THEAD\)|\(TBODY\)|\(TR\)|\(BR/\)', '', content)

            # elimina (DIV )(TABLE )
            content = re.sub(r'\(DIV style="color: rgb\(0, 0, 0\);font-family: &amp;quot;font-size: 14px;font-style: normal;font-weight: 400;letter-spacing: normal;text-align: start;text-indent: 0px;text-transform: none;white-space: normal;word-spacing: 0px"\)|\(TABLE style="border-collapse: collapse;border-spacing: 0px;margin: 1rem 0px 0px;padding: 0px;border: 0px;width: 860px;table-layout: auto;font-size: 0.875rem;color: rgb\(23, 23, 23\);font-family: &amp;quot"\)', '', content)

            content = re.sub(r'\(/TBODY\)|\(/TABLE\)|\(/DIV\)', '', content)

            content = re.sub(r'\(CODE style=".*?"\)\(SPAN style=".*?"\)(.*?)\(/SPAN\)\(/CODE\)', r'``\\1``', content)

            content = re.sub(r'\(CODE style=".*?"\)|\(/CODE\)', '`', content)

            content = re.sub(r'\((/?SPAN)\)', '', content, flags=re.IGNORECASE)

            # Anidar y aplicar ambas expresiones regulares en una sola línea
            content = re.sub(r'\|>', '|', re.sub(r'(\|[^|]*)(?=>)', r'\\1|', content))

            # Añade saltos de linea después de un pipe si se ve bloque de código
            content = re.sub(r'(\|\s*)\\n(```)', r'\\1\\n\\n\\2', content)

             
            # expresión regular para corregir los encabezados de la tabla Markdown       
            content = re.sub(r'(\|\s*\w+\s*\|\s*\w+\s*\|)\s*(\|\-\-\|\-\-\|)', r'\\1\\n\\2', content)

            # Expresión regular para detectar las filas de la tabla y ajustar los separadores
            content = re.sub(r'(\|\s*[^|]+\s*\|\s*[^|]+\s*\|\s*[^|]+\s*\|\s*[^|]+\s*\|)\s*\|--\|--\|--\|--\|--\|', r'\\1\\n|--|--|--|--|', content)

            # Modificamos la expresión regular fix tabla y header encabezado
            content = re.sub(r'\|\s*([^\\n|]+)\s*\\n\s*\|\s*([^\\n|]+)\s*\|\s*\|\s*---\s*---\s*\|', r'| \\1 | \\2 |\\n| --- | --- |', content)

            # Eliminar porcentaje si va acompañado de guión bajo
            content = re.sub(r'%_', '_', content)

            # Ajustes de etiquetas HTML espacios innecesarios
            content = re.sub(r'\s*((</?(?:b|span[^>]*)>))\s*', r'\\1', content)

            # Expresión regular para modificar la tabla en una sola línea
            content = re.sub(r'(\| Purpose User  \|)|(\|--\|--\|--\|)', lambda m: '| Purpose | User  |' if m.group(1) else '|--|--|--|--|', content)

            # (b)(span style="color:red")IMPORTANT:(/b)| realiza un salto de párrafo ya que no tiene tuberia al principio
            content = re.sub(r'^([^|]+)(\|)', r'\\1\\n\\2', content, flags=re.MULTILINE)

            # Paso 2 actualizado: Asegurarse de que las celdas vacías antes del encabezado tienen un salto de línea
            content = re.sub(r'(\|[ \t]*\|[ \t]*\|[ \t]*\|)(--\|--\|--\|--\|)', r'\\1\\n|\\2', content)

            # paso3 identifica https y http luego quita acentos graves innecesarios  
            content = re.sub(r'`(https?://[^\s`]+)`', r'\\1', content)
            
            # Cuando la etiqueta (b)span style está pega a dos puntos
            content = re.sub(r':(?=\(b\)|\(span style="color:[^"]+"\))', ':\\n', content)
            
            # Expresión regular que añade `(/span)` si falta después de `(/b)`, con cualquier color en el estilo
            content = re.sub(r"(\(b\)\(span style=\\"color:[a-zA-Z]+\\"\).*?\(\/b\))(?!\(\/span\))", r"\\1(/span)", content)
            
            # Para espacios dobles después de la etiqueta de apertura
            content = re.sub(r'(\(span [^)]+\).*?)(\s{2,})', r'\\1(/span)\\2', content)  # Para espacios dobles
  
            # Para finales de línea sin (/span) agrega 1 (/span)Para finales de línea sin (/span)
            content = re.sub(r'\((span [^)]+)\)([^(/\\n]+?)(?<!(/span))(?=\s*\||(/b)|(\\n)|$)', r'(\\1)\\2 (/span)', content) 
  
            # Expresión regular para buscar `(/b)(/span)(/span)` y eliminar solo el último `(/span)`
            content = re.sub(r"\(/b\)\(/span\)\(/span\)", r"(/b)(/span)", content)
            
            # Expresión regular para quitar '(/span)' que está cerca de una tubería '|'
            content = re.sub(r'\|(\s*)\(/span\)', r'|', content)
            
            # Usamos re.sub para agregar un salto de línea después de los dos puntos y antes de "negrit"
            content = re.sub(r":(negrit````)", r":\\n\\1", content)
            
            # Usamos re.sub para agregar un salto de línea si detrás de los dos puntos hay palabras y luego tres guiones
            content = re.sub(r"(\w+:\s*)---", r"\\1\\n---", content)
            
            # Expresión regular para quitar '(/span)' que está cerca de un '.com'
            content = re.sub(r'(\.com)\s*\(/span\)', r'\\1', content)
            
            # Usar re.sub para eliminar los saltos de línea intermedios SOLO cuando hay 3 guiones
            content = re.sub(r'\\n{2,}(?=---)', '\\n', content)
            
            # Insertar saltos de párrafos después de los dos puntos antes de la imagen HTML
            content = re.sub(r':(?!\\n)(?=\(a href="[^"]+"\)\(img src="[^"]+"[^>]*\/\)\(\/a\))', ':\\n', content)
            
            
            # Genera 1 espacio de más después de los 2 puntos
            content = re.sub(r'(?<!`\*\*):(?!\s|\*)(?![a-zA-Z0-9]*\*\*`)(?=[A-Za-z0-9])', ': ', content)

            # inserta un salto de linea en el caso que no haya ningún espacio entre los dos puntos y el texto 
            content = re.sub(r'(:)(?=(?!\*\*)(?:[\w`]|```))', r'\\1\\n', content)

            # Funciona todo alinea a la misma altura un asterisco con otro asterisco
            content = re.sub(r'(^.*:\s*\\n)(\*)(\s*\S.*)', r'\\1    \\2\\3', content, flags=re.MULTILINE)

            # Esta expresión regular quitar inserción de saltos de linea innecesarios dentro bloques de código márkdowns
            content = re.sub(r'```(.*?)\\n(.*?)\\n([ \\t]*\\|.*?)```', r'```\\1\\n\\2 \\3```', content, flags=re.DOTALL)
            
            
            # expresión regular que autoajusta espacios necesarios cabeceras de las tablas markdown
            # Función de reemplazo
            def replacer(m):
                # Si coincidió la parte separadora (?P<sep>)
                if m.group("sep") is not None:
                    # Convertir la línea separadora al formato "|--|--|..." según cantidad de columnas
                    sep_line = m.group("sep").strip()
                    columns = [col for col in sep_line.split("|") if col.strip()]
                    return "|" + "|".join(["--"] * len(columns)) + "|"
                else:
                    # Sino es la parte de encabezado (?P<header>), quitamos espacios iniciales
                    return m.group("header").lstrip()

            # Expresión regular aplicada directamente con re.sub()
            content = re.sub(r"^(?P<sep>\s*\|(?:\s*[-:]+\s*\|)+\s*)$|^(?P<header>\s+\|(?:[^|]*\|)+.*)$", 
                                        replacer, 
                                        content, 
                                        flags=re.MULTILINE)
                            
  
            
            # Ajusta la r la pone de bajo en el código YAML 
            content = re.sub(r'(```[\s\S]*?```|\(code\)[\s\S]*?\(/code\))', lambda m: re.sub(r'(^[ \t]*-[ \t]*script:[ \t]*\\|)[ \t]*(#.*$)', lambda a: f"{a.group(1)}\\n{' ' * a.group(1).find('r')}{a.group(2).strip()}", m.group(0), flags=re.MULTILINE), content, flags=re.DOTALL)

            # 1) Aplana pipes dentro de (code)…(/code) sin tocar los comentarios
            content = re.sub(r'(```[\s\S]*?```|\(code\)[\s\S]*?\(/code\))', lambda m: m.group(0).replace("\\n|", " |").replace("| \\n", " | "), content, flags=re.DOTALL)

            content = re.sub(
                r'(```[\s\S]*?```|\(code\)[\s\S]*?\(/code\))', 
                lambda m: re.sub(
                    r'(-)[ \t]+(script)[ \t]*(:)[ \t]*(?=\\|)',
                    r'\\1 \\2\\3 ',
                    m.group(0)
                ), 
                content, 
                flags=re.DOTALL
            )

            content = re.sub(r'(```[\s\S]*?```|\(code\)[\s\S]*?\(/code\))', lambda m: re.sub(r'(^[ \t]*-[ \t]*script:[ \t]*\|)[ \t]*(#.*$)', lambda a: f"{a.group(1)}\\n{' ' * a.group(1).find('r')}{a.group(2).strip()}", m.group(0), flags=re.MULTILINE), content, flags=re.DOTALL)

            # Solamente calcula y  convierte de un signo > a parentesis blockquote. Convierte todos los blockquote borde izquierdo en parentesis para que se pueda procesar bien
            lines = content.strip('\\n').splitlines()
            result = []
            prev_blockquote = False

            for line in lines:
                if re.match(r'^\s*>+\s*(.*)', line):
                    result.append(re.sub(r'^\s*>+\s*(.*)', r') \\1', line))
                    prev_blockquote = True
                elif re.match(r'^\s*\|\s*.*', line) and prev_blockquote:
                    result.append(re.sub(r'^\s*\|\s*(.*)', r') \\1', line))
                    prev_blockquote = False  # solo permitir uno seguido
                else:
                    result.append(line)
                    prev_blockquote = False  # se corta la cadena

            content = '\\n'.join(result)          

            info = {
                'name': page['path'],
                'short_name': page['path'].split('/')[-1],
                'url': page['url'],
                'original_content': original_content,
                'content': content,
                'level': level,
                'subpages': []
            }

            page_info = [info]

            if 'subPages' in page:
                for sub_page in page['subPages']:
                    page_info.extend(extract_pages_recursive(sub_page, headers, level+1))

            return page_info

        def extract_url_values(url):
            regex = r"https://dev.azure.com/(?P<organization>[^/]+)/(?P<project>[^/]+)/_wiki/wikis/(?P<wiki>[^/]+)/(?P<page_id>\d+)/.*"
            match = re.search(regex, url)
            if match:
                return match.groupdict()
            else:
                return None

      
        def obtain_page(organization, project, wiki, page_id, headers):
            api_url = f"https://dev.azure.com/{organization}/{project}/_apis/wiki/wikis/{wiki}/pages/{page_id}?api-version=7.0&recursionLevel=full&includeContent=true"
            response = requests.get(api_url, headers=headers)
            if response.status_code == 200:
                print("Successful response http 200.")
                page = json.loads(response.text)
                return page
            else:
                print(f"Error al obtener la página de la Wiki: {response.status_code}")

            return None
    

        def download_specific_page(headers, wiki_url2):
            url_values = extract_url_values(wiki_url2)

            if url_values is None:
                print("La URL proporcionada no es válida.")
                return

            page_id = url_values['page_id']
            organization = url_values['organization']
            project = url_values['project']
            wiki = url_values['wiki']

            page = obtain_page(organization, project, wiki, page_id, headers)

            # Añade una condición de parada para la recursión
            if page is not None:
                return organization, project, wiki, page
            else:
                print("No se pudo descargar la página.")
                return None

              
        # Pedir la información al usuario una sola vez y almacenarla en variables globales
        # Verificar si las credenciales ya se han almacenado
        if stored_wiki_url is None or stored_personal_access_token is None or stored_wiki_url2 is None:
            wiki_url2 = input("Introduce la URL de la página que quieres buscar: ")
            personal_access_token = input("Introduce tu token de acceso personal: ")
            print("Por favor usuario necesito que especifiques tu URL principal azure wiki para poder descargar los ficheros tanto fotos u otros ficheros")
            wiki_url = input("Introduce la URL principal del portal Wiki de Azure: ")
            # Aquí iría la lógica para verificar si las credenciales son válidas.
            # Si son válidas, las almacenamos en las variables globales
            stored_wiki_url = wiki_url
            stored_personal_access_token = personal_access_token
        else:
            # Usar las credenciales almacenadas
            wiki_url = stored_wiki_url
            personal_access_token = stored_personal_access_token

        
        credentials = f":{personal_access_token}"
        encoded_credentials = base64.b64encode(credentials.encode('utf-8')).decode('utf-8')
        headers = {
                'Authorization': f'Basic {encoded_credentials}',
                "Content-Type": "application/json",
                "Accept": "application/json"
        }

        result = download_specific_page(headers, wiki_url2)
        if result is not None:
              organization, project, wiki, page_data = result
              page_info = extract_pages_recursive(page_data, headers)  # Asegúrate de que esto devuelva una lista 
        else:
            print("No se pudo descargar la página.")
            

        def extract_url_values(wiki_url2):
            # Decodifica la URL para manejar caracteres especiales
            decoded_wiki_url2 = urllib.parse.unquote(wiki_url2)
            # Usa regex para extraer la organización, el proyecto y el wiki de la URL decodificada
            api_url_match = re.search(r"https://dev.azure.com/(?P<organization>[^/]+)/(?P<project>[^/]+)/_wiki/wikis/(?P<wiki>.+).wiki/(?P<page_id>\d+)/.*", decoded_wiki_url2)
            if api_url_match:
                 organization = api_url_match.group('organization')
                 project = api_url_match.group('project')
                 wiki = api_url_match.group('wiki')
                 page_id = api_url_match.group('page_id')
                 return {'organization': organization, 'project': project, 'wiki': wiki, 'page_id': page_id}
            else:
                 print("La URL proporcionada no es válida.")
                 exit() 

        
        # Escribir el contenido original de las páginas en un archivo Markdown
        with open('htmlymd.md', 'w', encoding='utf-8') as f:
            for page in page_info:
                # Obtener el título y contenido original de la página
                title = page['name'].split('/')[-1]
                original_content = page['original_content'].strip()

                # Solo escribir el título y contenido en el archivo Markdown si el contenido no está vacío
                if original_content:
                    f.write(f'# {title}\\n')
                    f.write(original_content)
                    f.write('\\n\\n')

                # Recorrer las subpáginas y escribir su contenido original
                for subpage in page['subpages']:
                    # Obtener el título y contenido original de la subpágina
                    subpage_title = subpage['name'].split('/')[-1]
                    subpage_original_content = subpage['original_content'].strip()

                    # Solo escribir el título y contenido en el archivo Markdown si el contenido no está vacío
                    if subpage_original_content:
                        f.write(f'# {subpage_title}\\n')
                        f.write(subpage_original_content)
                        f.write('\\n\\n')
                
        # Escribir el contenido de las páginas en un archivo Markdown
        with open('todosmd.md', 'w', encoding='utf-8') as f:
            for page in page_info:
                # Obtener el título y contenido de la página
                title = page['name'].split('/')[-1]
                content = page['content'].strip()

                # Solo escribir el título y contenido en el archivo Markdown si el contenido no está vacío
                if content: 
                    f.write(f'# {title}\\n')
                    f.write(content)
                    f.write('\\n\\n')

                # Recorrer las subpáginas y escribir su contenido
                for subpage in page['subpages']:
                    # Obtener el título y contenido de la subpágina
                    subpage_title = subpage['name'].split('/')[-1]
                    subpage_content = subpage['content'].strip()

                    # Solo escribir el título y contenido en el archivo Markdown si el contenido no está vacío
                    if subpage_content:
                        f.write(f'# {subpage_title}\\n')
                        f.write(subpage_content)
                        f.write('\\n\\n')
                        
                else:
                    if not page['content'].strip():
                            print(f"No se encontró la página con ID {page_id}.")


        def extract_placeholders(template_path):
            doc = Document(template_path)
            placeholders = []

            previous_full_text = None  # Agregar variable para almacenar el texto completo anterior

            for element in doc.element.body.iter():
                if element.tag.endswith('p'):
                    paragraph = Paragraph(element, doc)
                    if hasattr(paragraph, 'runs'):
                        # Concatenate the text of adjacent runs
                        full_text = ''.join([run.text for run in paragraph.runs])
                        if full_text != previous_full_text:  # Solo imprimir el texto completo si es diferente del anterior
                            previous_full_text = full_text  # Actualizar el valor de previous_full_text
                        
                        # Find placeholders in the full text
                        start_index = full_text.find('{{')
                        end_index = full_text.find('}}')
                        while start_index != -1 and end_index != -1:
                            placeholder = full_text[start_index+2:end_index].strip()
                            placeholders.append(placeholder)
                            
                            # Find the next placeholder
                            start_index = full_text.find('{{', end_index)
                            end_index = full_text.find('}}', end_index+2)

                elif element.tag.endswith('tc'):  # Check for table cell (td)
                    for p in element.iterchildren('{%s}p' % nsmap['w']):
                        paragraph = Paragraph(p, doc)
                        if hasattr(paragraph, 'runs'):
                            # Concatenate the text of adjacent runs
                            full_text = ''.join([run.text for run in paragraph.runs])
                            if full_text != previous_full_text:  # Solo imprimir el texto completo si es diferente del anterior
                                previous_full_text = full_text  # Actualizar el valor de previous_full_text
                            
                            # Find placeholders in the full text
                            start_index = full_text.find('{{')
                            end_index = full_text.find('}}')
                            while start_index != -1 and end_index != -1:
                                placeholder = full_text[start_index+2:end_index].strip()
                                placeholders.append(placeholder)
                                
                                # Find the next placeholder
                                start_index = full_text.find('{{', end_index)
                                end_index = full_text.find('}}', end_index+2)

            # Remove duplicates from the list of placeholders while maintaining the order of the elements
            placeholders = list(OrderedDict.fromkeys(placeholders))

            return placeholders

        def create_context(page_info, placeholders):
            context = {}
            title_index = 0
            for placeholder in placeholders:
                if title_index < len(page_info):
                    page = page_info[title_index]
                    # Agregar valores al contexto tanto para marcadores de posición de título como para marcadores de posición de contenido
                    if placeholder.endswith('_content'):
                        # Eliminar espacios adicionales del contenido
                        content = re.sub(r'\s+', ' ', page['content'])
                        # Eliminar las dos almohadillas del contenido
                        content = content.replace('##', '')
                        context[placeholder] = page['content']
                        title_index += 1
                    else:
                        context[placeholder] = page['name'].split('/')[-1]
            return context


        def add_titles_to_template(template_path, titles_info):
            doc = Document(template_path)

            # Find the element after which to insert the placeholders
            insert_element = None
            for element in doc.element.body.iter():
                if element.tag.endswith('p'):
                    paragraph = Paragraph(element, doc)
                    if hasattr(paragraph, 'runs'):
                        # Concatenate the text of adjacent runs
                        full_text = ''.join([run.text for run in paragraph.runs])
                        if 'REFERENCES' in full_text:
                            insert_element = paragraph
                            break

            if insert_element is not None:
                # Insert the titles before the element with the desired heading style
                for title_info in titles_info:
                    sanitized_title = sanitize_placeholder(title_info["title"])
                    p = insert_element.insert_paragraph_before(f'{{{{{sanitized_title}}}}}')
                    # Map levels to Word Heading styles
                    level_map = {1: 'Heading 1', 2: 'Heading 2', 3: 'Heading 3', 4: 'Heading 4', 5: 'Heading 5'}
                    heading_style = level_map.get(title_info["level"], 'Normal')
                    p.style = heading_style

                    # Insert a placeholder for the content below the title
                    p = insert_element.insert_paragraph_before(f'{{{{{sanitized_title}_content}}}}')
                    # Set the alignment of the paragraph to left
                    p.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.LEFT
                    
                    # Apply bold formatting to the text if it starts with two hash symbols and remove the hash symbols from the content
                    if p.text.startswith('##'):
                        for run in p.runs:
                            run.bold = True
                            run.text = run.text.replace('##', '')

                # Insert a page break after the last title
                p = insert_element.insert_paragraph_before()
                run = p.add_run()
                run.add_break(docx.enum.text.WD_BREAK.PAGE)

            # Save the modified template
            doc.save('new_template.docx')


        # Generate the title information
        titles_info = [{'title': page['name'].split('/')[-1], 'level': page['level']} for page in page_info if page['content'].strip() and page['name'].split('/')[-1].strip()]


        # Call the function to create a new template and add the titles to it
        add_titles_to_template('plantilla.docx', titles_info)

        # Load the modified template
        doc = DocxTemplate("new_template.docx")

        # Extract the placeholders from the template
        placeholders = extract_placeholders("new_template.docx")

        # Create a new context
        context = create_context(page_info, placeholders)

        # Render template with dynamic context
        doc.render(context)

        # Save the generated document
        doc.save("documento_generado.docx")


        print("Wait one moment to finish work")
        # Wait for a few seconds to make sure the file is saved
        time.sleep(4)
        print("already")


        # Check if the file was saved successfully   
        if os.path.exists("documento_generado.docx"):
            print("The documento_generado.docx file has been saved successfully.")
        else:
            print("No se pudo guardar el archivo documento_generado.docx.")

        # Get the directory of the script
        dir_path = os.path.dirname(os.path.realpath(__file__))

        # Build the file path
        docx_file = os.path.join(dir_path, "documento_generado.docx")

        # Call the function to update the table of contents in the generated document
        update_toc(docx_file)

        document = Document('documento_generado.docx')

        def clean_attachments_folder():
            ###Elimina todos los archivos en la carpeta .attachments
            attachments_dir = '.attachments'
            if os.path.exists(attachments_dir):
                for file_name in os.listdir(attachments_dir):
                    file_path = os.path.join(attachments_dir, file_name)
                    try:
                        if os.path.isfile(file_path) or os.path.islink(file_path):
                            os.unlink(file_path)
                            print(f"Deleted file: {file_name}")
                    except Exception as e:
                        print(f"Failed to delete {file_name}: {e}")
            else:
                os.makedirs(attachments_dir)  # Si no existe, la creamos
        # Limpiar la carpeta .attachments antes de descargar nuevos archivos
        clean_attachments_folder()

        print("Preparandose para descargar azure wiki GIT")

        print("Se encarga esta función en descargar de forma recursiva los ficheros de GIT azure")
        list_url = f"https://dev.azure.com/{organization}/{project}/_apis/git/repositories/{wiki}/items?scopePath=/.attachments&recursionLevel=full&api-version=5.0"

        # Configura la autenticación
        headers = {
            'Authorization': f'Basic {base64.b64encode((":{}".format(personal_access_token)).encode()).decode()}'
        }

        # Crea la carpeta .attachments si no existe
        if not os.path.exists('.attachments'):
            os.makedirs('.attachments')

        # Realiza la solicitud GET para obtener la lista de archivos
        list_response = requests.get(list_url, headers=headers)

        # Procesa la respuesta
        if list_response.status_code == 200:
            response_json = list_response.json()
            if 'value' in response_json:
                files = response_json['value']
                folders = [file for file in files if file.get('isFolder')]
                for folder in folders:
                    print(f"Your folder is {folder['path']}")
                for file in files:
                    # Salta el elemento que representa la carpeta en sí
                    if file.get('isFolder'):
                        continue

                    # Descarga cada archivo
                    download_url = f"https://dev.azure.com/{organization}/{project}/_apis/git/repositories/{wiki}/items?path={file['path']}&api-version=5.0"
                    download_response = requests.get(download_url, headers=headers)
                    if download_response.status_code == 200:
                        file_name = os.path.basename(file['path'])
                        with open(f'.attachments/{file_name}', 'wb') as f:
                            f.write(download_response.content)
                        print(f"Successfully downloaded the file {file_name} de la carpeta {folder['path']}")
                    else:
                        print(f"Error al descargar {file['path']} from the folder {folder['path']}: {download_response.status_code}")
            else:
                print("La clave 'value' no está presente en la respuesta de la API.")
        else:
            print(f"Error al listar archivos: {list_response.status_code}. Ese código 404 no te preocupues simplemente en tu wiki no tienes ficheros imágenes de momento")


    paginaconcreta2()  # Llama a la función paginaconcreta2


# Definir las constantes manualmente
wdStatisticPages = 2
wdGoToPage = 1
wdGoToAbsolute = 1


print("Starting the Word application and opening the document...")
        
try:
    # Initialize Word application and open the document
    word_app = win32com.client.gencache.EnsureDispatch('Word.Application')
    word_app.DisplayAlerts = False
    word_app.Visible = False
    # Establecer Word en modo sin interfaz gráfica (evitar cualquier ventana emergente)
    word_app.ScreenUpdating = False
    word_app.DisplayAlerts = 0  # 0 = No alerts, 1 = alerts (por si acaso)
except Exception as e:
    print(f"An error occurred while trying to start the Word application: {e}")
    exit(1)

try:
    file_path = os.path.abspath("documento_generado.docx")
    doc = word_app.Documents.Open(file_path)
except Exception as e:
    print(f"An error occurred while trying to open the document: {e}")
    word_app.Quit()
    exit(1)


# Reemplazo de "↵" por "^p"
find_object = doc.Content.Find
find_object.ClearFormatting()
find_object.Text = '^l'
find_object.Replacement.ClearFormatting()
find_object.Replacement.Text = '^p'
find_object.Execute(Replace=2)




MONO_FONTS  = {"Consolas", "Courier New", "Courier", "Menlo", "Monaco"}
code_starts = set()

# ——————————————————————————————
# 1) Detectar párrafos de “código”
# ——————————————————————————————
in_fence = False
for para in doc.Paragraphs:
    txt = para.Range.Text.strip()

    # 1.1) Fence Markdown 
    if txt.startswith("```"):
        code_starts.add(para.Range.Start)
        in_fence = not in_fence
        continue
    if in_fence:
        code_starts.add(para.Range.Start)
        continue

    # 1.2) Fence de empresa (code)…(/code) — cláusula compacta
    if txt.startswith("(code)") or txt.startswith("(/code)"):
        code_starts.add(para.Range.Start)
        in_fence = txt.startswith("(code)")
        continue
    if in_fence:
        code_starts.add(para.Range.Start)
        continue

    # 1.3) Sombreado/fuente monoespaciada
    try:
        if para.Range.Shading.BackgroundPatternColor != const.wdColorAutomatic:
            code_starts.add(para.Range.Start)
            continue
    except:
        pass

    try:
        if para.Range.Font.Name in MONO_FONTS:
            code_starts.add(para.Range.Start)
    except:
        pass

def is_codeblock(para):
    return para.Range.Start in code_starts

# ——————————————————————————————
# 2) Función de reemplazo regex y aplicación formato
#    + Limpieza de numeración automática si comienza por número tras el reemplazo
# ——————————————————————————————
def regex_replace_and_format(doc, pattern, *, bold=False, italic=False, size=None):
    regex = re.compile(pattern)
    in_fence = False

    for para in doc.Paragraphs:
        txt = para.Range.Text.strip()

        # 2.1) Fence Markdown
        if txt.startswith("```"):
            in_fence = not in_fence
            continue

        # 2.2) Fence de empresa (code)…(/code)
        if txt.startswith("(code)") or txt.startswith("(/code)"):
            in_fence = txt.startswith("(code)")
            continue

        # Saltar dentro de cualquier bloque de código
        if in_fence or is_codeblock(para):
            continue

        # Aplicar reemplazo y formato
        text   = para.Range.Text
        offset = 0
        for m in regex.finditer(text):
            s, e    = m.span()
            start   = para.Range.Start + s + offset
            end     = para.Range.Start + e + offset
            rng     = doc.Range(Start=start, End=end)

            rng.Text = m.group(1)

            # Limpia autolista numerada tras reemplazo si empieza por número
            if re.match(r'^\s*\d+\.\s', rng.Text):
                try:
                    rng.ListFormat.RemoveNumbers()
                except:
                    pass

            if bold:   rng.Bold      = True
            if italic: rng.Italic    = True
            if size:   rng.Font.Size = size

            offset -= (len(m.group(0)) - len(m.group(1)))

# ——————————————————————————————
# 3) Lógica de reemplazo Markdown
# ——————————————————————————————
# Negrita **texto**
regex_replace_and_format(doc, r'\*\*(.+?)\*\*', bold=True)

# Cursiva *texto* y _texto_
regex_replace_and_format(doc, r'\*(.+?)\*', italic=True)
regex_replace_and_format(doc, r'_(.+?)_', italic=True)

# Títulos con # (de mayor a menor profundidad)
regex_replace_and_format(doc, r'#####\s*(.+)', bold=True, size=11)
regex_replace_and_format(doc, r'####\s*(.+)',  bold=True, size=11)
regex_replace_and_format(doc, r'###\s*(.+)',   bold=True, size=11)
regex_replace_and_format(doc, r'##\s*(.+)',    bold=True, size=16)
regex_replace_and_format(doc, r'#\s*(.+)',     bold=True, size=16)

# ——————————————————————————————
# 4) Limpieza final de autolistas accidentales
# ——————————————————————————————
for para in doc.Paragraphs:
    txt = para.Range.Text.strip()
    if re.match(r'^\d+\.\s', txt) and para.Range.Font.Bold:
        try:
            para.Range.ListFormat.RemoveNumbers()
        except:
            pass


           
print("Bullet and sub ul li HTML robot tags only matched")

opt = doc.Application.Options

# --- AutoFormat As You Type (al escribir) ---
opt.AutoFormatAsYouTypeApplyNumberedLists      = False
opt.AutoFormatAsYouTypeApplyBulletedLists      = False
opt.AutoFormatAsYouTypeFormatListItemBeginning = False
opt.AutoFormatAsYouTypeApplyHeadings           = False

# --- AutoFormat (al formatear rangos/documento) ---
opt.AutoFormatApplyLists         = False
opt.AutoFormatApplyBulletedLists = False
opt.AutoFormatApplyHeadings      = False


en_lista = False
en_bloque_codigo = False

texto_completo = doc.Range().Text
etiquetas_ul_suelta = '(ul)' in texto_completo and '(/ul)' not in texto_completo
etiquetas_li_suelta = '(li)' in texto_completo and '(/li)' not in texto_completo

def es_encabezado_markdown(paragraph: str) -> bool:
    return bool(re.match(r'^\s*#+\s*', paragraph))

for par in doc.Paragraphs:
    paragraph = par.Range.Text
    stripped = paragraph.strip()

    # 1) detectar fenced code Markdown y tus tags de empresa
    if stripped.startswith("```"):
        en_bloque_codigo = not en_bloque_codigo
        continue
    # bloque de código empresa: (code) … (/code)
    if stripped.startswith("(code)") or stripped.startswith("(/code)"):
        en_bloque_codigo = stripped.startswith("(code)")
        continue
    # Mientras estemos dentro de cualquier bloque de código, saltamos
    if en_bloque_codigo:
        continue

    # 1.b) si es encabezado Markdown, quitamos autolist
    if es_encabezado_markdown(paragraph):
        par.Range.ListFormat.RemoveNumbers()
        continue

    # 2) detectar encabezados numerados “1. ”, “2. ”…
    if re.match(r'^\s*\d+\.\s', paragraph):
        par.Range.ListFormat.RemoveNumbers()
        continue

    # 3) manejar tus etiquetas (ul)/(li)
    if etiquetas_ul_suelta and stripped == "(ul)":
        par.Range.Delete()
        continue
    if etiquetas_li_suelta and stripped == "(li)":
        par.Range.Delete()
        continue
    if stripped == "(ul)":
        en_lista = True
        par.Range.Delete()
        continue
    elif stripped == "(/ul)":
        en_lista = False
        par.Range.Delete()
        continue

    # 4) bullets de nivel 2 (dos espacios + guión)
    if paragraph.startswith('  - '):
        par.Format.SpaceAfter = 0
        par.Range.ListFormat.ApplyListTemplateWithLevel(
            ListTemplate=par.Range.Application.ListGalleries.Item(1).ListTemplates.Item(1),
            ContinuePreviousList=True,
            ApplyTo=constants.wdListApplyToWholeList,
            DefaultListBehavior=win32.constants.wdWord10ListBehavior
        )
        par.Range.ListFormat.ListLevelNumber = 2
        par.Range.ListFormat.ListTemplate.ListLevels(2).NumberFormat = chr(9675)
        continue

    # 5) bullets de nivel 1 (lista principal)
    if en_lista or paragraph.lstrip().startswith('- '):
        par.Format.SpaceAfter = 0
        par.Range.ListFormat.ApplyListTemplateWithLevel(
            ListTemplate=par.Range.Application.ListGalleries.Item(1).ListTemplates.Item(1),
            ContinuePreviousList=True,
            ApplyTo=constants.wdListApplyToWholeList,
            DefaultListBehavior=win32.constants.wdWord10ListBehavior
        )
        par.Range.ListFormat.ListLevelNumber = 1
        continue




print("table creation type 2 the newest")
# Lista para almacenar todas las tablas encontradas
all_tables_data = []

# Patrón regex para encontrar enlaces en Markdown
pattern = re.compile(r'\[([^\]]+)\]\(([^)]+)\)')

# Patrón regex para encontrar imágenes en Markdown rodeadas por tuberías
image1_pattern1 = re.compile(r'\|\!\[([^\]]*)\]\((.*?)\)\|')

table = None  # Añade esta línea para inicializar 'table' como None


def add_missing_delimiters(table_lines):
    new_table_lines = []
    for line in table_lines:
        stripped_line = line.rstrip()  # Elimina los espacios en blanco al final
        if stripped_line and not stripped_line.endswith("|"):
            # Si la línea no está vacía y no termina con '|', añade '|'
            stripped_line += " |"
        new_table_lines.append(stripped_line)
    return new_table_lines


while True:
        try:    
            start_table = None
            end_table = None
            table_lines = []
            blank_line_count = 0  # Move this outside the loop

            # Búsqueda de tablas en el documento
            for index, para in enumerate(doc.Paragraphs):
                line = para.Range.Text.strip()
                placeholders = []  # Definir 'placeholders' si es necesario

                # Si la línea contiene una imagen en formato Markdown rodeadas por tuberías, ignórala
                if image1_pattern1.search(line):
                    continue

                # Ignora como encabezados pipelines dentro de acentos invertidos bloque de código
                if re.search(r'`+[^`]*\|[^`]*`+|".*?\|.*?"|\|\|', line):
                    continue
 
 
                # Ignora cualquier bloque (a …)(img …)(/a)
                # Ignorar:
                # 1) Cualquier bloque tipo (a …)(img …)(/a)
                # 2) Cualquier fila con solo una URL de photo-pick.com/online/... .link (falso encabezado)
                if (
                    ("(img" in line and re.search(r"\(a\s+href=", line, re.IGNORECASE))
                    or re.fullmatch(
                        r"\|\s*https?://(?:www\.)?photo-pick\.com/online/[^\s|]+\.link\s*\|",
                        line.strip(),
                        re.IGNORECASE,
                    )
                ):
                    continue

                # Ignorar líneas que comienzan con una tubería seguida de '+-' o '-+'
                if re.match(r'^\|\s*\+-', line) or re.match(r'^\|\s*-+\+', line):
                    continue

                #  Ignorar filas tipo |(img …)|
                if re.fullmatch(r'^\|\(img\s+[^)]*\)\|$', line, re.IGNORECASE):
                    continue
    
                # Manejar inicio y fin de tabla
                if "|" in line and line.lstrip().startswith("|") and len(line.split("|")) > 2:
                    if start_table is None:
                        start_table = index
                    blank_line_count = 0  # Reiniciar contador de líneas en blanco
                    table_lines.append(line.strip())
                else:
                    if start_table is not None:
                        blank_line_count += 1
                        if blank_line_count > 1:
                            end_table = index
                            break
                    continue


            # Llama a la función para añadir delimitadores faltantes
            table_lines = add_missing_delimiters(table_lines)

            # Si no se encuentra ninguna tabla, se detiene el bucle
            if not table_lines:
                break

            # Procesamiento de la tabla en markdown
            data = []
            headers_found = False
            header_separator_line_index = None  # Índice de la línea de separación del encabezado
            for line in table_lines:
                if not headers_found and "|" in line:
                    # Procesar los encabezados
                    headers = [cell.strip() for cell in line.split('|')[1:-1]]
                    headers = [re.sub(r':?----+:?', '', header).strip() for header in headers]  # Eliminar patrones de alineación
                    if headers[-1].strip() == '':
                        headers = headers[:-1]
                    print(f"Encabezados procesados en esta línea: {headers}")  # Imprime los encabezados 
                    data.append(headers)
                    headers_found = True
                elif headers_found and (re.match(r'^\|\s*-+\s*\|', line) or re.match(r'^\|\s*(:?----+:?\s*\|)+', line)):
                    # Ignora la línea de separación del encabezado, no se añade a 'data'
                    continue
                elif "|" in line:
                    # Procesar las celdas de datos
                    cells = [cell.strip() for cell in line.split('|')[1:-1]]
                    if cells[-1].strip() == '':
                        cells = cells[:-1]
                    cells = [cell.replace(':white_check_mark:', '✅') for cell in cells]
                    data.append(cells)

            # Eliminar la fila de separación de encabezados si existe
            if header_separator_line_index is not None:
                del data[header_separator_line_index]

            # Validar si los datos de la tabla están completos
            if len(data) < 2 or len(data[0]) < 1:
                print(f"Los datos de la tabla están incompletos: {data}")
                continue  # Saltar al siguiente ciclo si no hay datos válidos

            # Ajustar el rango para que solo incluya la tabla
            start_range = None
            end_range = None
        
            counter = 0
            for para in doc.Paragraphs:
                counter += 1
                if counter == start_table:
                    start_range = para.Range.Start
                if counter == end_table:
                    end_range = para.Range.End
                    break

            # Evitar modificar un rango no válido
            if start_range is not None and end_range is not None and start_range < end_range:
                table_text = doc.Range(start_range, end_range).Text
                if "|" not in table_text:
                    print("El rango no contiene una tabla. Saltando...")
                    continue
            else:
                print(f"Rango no válido: start={start_range}, end={end_range}")
                continue
            
            doc.Range(start_range, end_range).Delete()

            table_range = doc.Range(start_range, start_range)

            cm_to_points = 2.15 * 28.3465  # 1 cm es aproximadamente 28.3465 puntos
            
            if len(data) > 2 and len(data[0]) > 0:
                table = doc.Tables.Add(table_range, len(data), len(data[0]))
                # Configura el ancho de cada celda de la tabla
                if table is not None:
                    for row in table.Rows:
                        for cell in row.Cells:
                            try:
                                cell.Width = cm_to_points
                            except Exception as e:
                                    print(f"Error al ajustar el ancho de la celda: {e}")
            else:
                print("The markdown table in this bot is not recognized.")
                continue  # Continúa con el siguiente ciclo del bucle while si la tabla no se reconoce

            markdown_link_found_in_table = False  # Añade esta línea para inicializar 'markdown_link_found_in_table' en False

            for i, row_data in enumerate(data):
                for j, cell_data in enumerate(row_data):
                    # Comprueba si los índices están dentro del rango de la tabla
                    if i < table.Rows.Count and j < table.Columns.Count:
                        cell = table.Cell(i+1, j+1)
                        cell_range = cell.Range             

                        # Intenta ajustar el formato del párrafo
                        cell_range.ParagraphFormat.ListTemplate = None
                        # Ajusta el estilo del párrafo a 'Normal'
                        cell_range.Style = doc.Styles('Normal')
                        
                        cell_range.Text = cell_data.strip()
        
                        # Manipulación de hipervínculos
                        matches = pattern.findall(cell_data)
                        
                        if matches:
                            markdown_link_found_in_table = True  # Se encontró un enlace Markdown en una celda de la tabla

                            for text, url in matches:
                                hyperlink_range = cell_range.Duplicate
                            
                                # Limpia el texto de la celda antes de añadir el hipervínculo.
                                cell_range.Text = text.strip()
                                
                                
                                hyperlink_range.Find.Execute(FindText=text)
                                doc.Hyperlinks.Add(Anchor=hyperlink_range, Address=url)
                                
                                
                    else:
                        print(f"Índice fuera de rango: i={i}, j={j}")

            table.Style = "Acc_Table_1"
            all_tables_data.append(data)
        except Exception as e:
                    print(f"Se produjo un error al procesar la tabla: {e}")
                    break



# Inicializa 'sheet_resized' como False
sheet_resized = False

# Recorrer las tablas y aumentar el tamaño de la hoja si hay más de 5 columnas
for table in doc.Tables:
    if table.Columns.Count > 5:
        try:
            # Aumentar el tamaño de la hoja en 2.54 cm (equivalente a 1 pulgada)
            points_in_cm = 5.45 * 28.3465  # 1 cm = 28.3465 puntos
            doc.PageSetup.PageWidth = doc.PageSetup.PageWidth + points_in_cm
            doc.PageSetup.PageHeight = doc.PageSetup.PageHeight + points_in_cm
            print(f"The size of the sheet has been increased. New width: {doc.PageSetup.PageWidth}, Nuevo alto: {doc.PageSetup.PageHeight}")
            sheet_resized = True
            break  # Salir del bucle después de la primera tabla encontrada
        except Exception as e:
            print(f"Se produjo un error al intentar ajustar el tamaño de la hoja: {e}")

if sheet_resized:
    print("The document sheet size has been resized.")
else:
    print("The size of the document sheet has not been altered.")



print("creating tables in word with pywin32 custom table size")
# Lista para almacenar todas las tablas encontradas
all_tables_data = []

# Patrón regex para encontrar enlaces en Markdown
pattern = re.compile(r'\[([^\]]+)\]\(([^)]+)\)')

# Patrón regex para encontrar imágenes en Markdown rodeadas por tuberías
image1_pattern1 = re.compile(r'\|\!\[([^\]]*)\]\((.*?)\)\|')

table = None  # Añade esta línea para inicializar 'table' como None

# Analizar cada párrafo buscando el inicio y el final de una tabla en markdown
while True:
    start_table = None
    end_table = None
    table_lines = []

    # Búsqueda de tablas en el documento
    for index, para in enumerate(doc.Paragraphs):
        line = para.Range.Text.strip()
        placeholders = []  # Definir 'placeholders' si es necesario
        
        # Verifica si la línea contiene una imagen en formato Markdown rodeada por tuberías
        if image1_pattern1.search(line):
            continue  # Si es así, ignora la línea y pasa a la siguiente

        # Ignora como encabezados pipelines dentro de acentos invertidos bloque de código
        if re.search(r'`+[^`]*\|[^`]*`+|".*?\|.*?"|\|\|', line):
            continue
                    
        # Ignora cualquier bloque (a …)(img …)(/a)
        # Ignorar:
        # 1) Cualquier bloque tipo (a …)(img …)(/a)
        # 2) Cualquier fila con solo una URL de photo-pick.com/online/... .link (falso encabezado)
        if (
            ("(img" in line and re.search(r"\(a\s+href=", line, re.IGNORECASE))
            or re.fullmatch(
                r"\|\s*https?://(?:www\.)?photo-pick\.com/online/[^\s|]+\.link\s*\|",
                line.strip(),
                re.IGNORECASE,
            )
        ):
            continue


        # Ignora las líneas que comienzan con una tubería seguida de cualquier número de espacios y luego '+-' o '-+'
        if re.match(r'^\|\s*\+-', line) or re.match(r'^\|\s*-+\+', line):
            continue
        
        #  Ignorar filas tipo |(img …)|
        if re.fullmatch(r'^\|\(img\s+[^)]*\)\|$', line, re.IGNORECASE):
            continue
                    
                    
        if "|" in line and line.lstrip().startswith("|") and len(line.split("|")) > 2:
            if start_table is None:
                start_table = index
            # Vuelve a insertar los enlaces originales (si fuera necesario)
            for placeholder, match in placeholders:
                line = line.replace(placeholder, match[0])
            table_lines.append(line.strip())
        elif start_table is not None:
            end_table = index
            break


    # Si no se encuentra ninguna tabla, se detiene el bucle
    if not table_lines:
        break

    # Procesamiento de la tabla en markdown
    data = []
    headers_found = False
    headers = []  # Añade esta línea para inicializar 'headers' como una lista vacía 
    for line in table_lines:
        if "|--" in line:
            headers_found = True
            continue
        if not headers_found:
            headers = [cell.strip() for cell in line.split('|')[1:]]
            if headers[-1].strip() == '':
                headers = headers[:-1]
            print(f"Encabezados procesados en esta línea: {headers}")  # Imprime los encabezados        
            continue
        else:
            cells = [cell.strip() for cell in line.split('|')[1:]]
            if cells[-1].strip() == '':
                cells = cells[:-1]
                cells = [cell.replace(':white_check_mark:', '✅') for cell in cells]
            print(f"Celdas procesadas en esta línea: {cells}")
            data.append(cells)

    data.insert(0, headers)

    counter = 0
    for para in doc.Paragraphs:
        counter += 1
        if counter == start_table:
            start_range = para.Range.Start
        if counter == end_table:
            end_range = para.Range.End
            break

    doc.Range(start_range, end_range).Delete()

    table_range = doc.Range(start_range, start_range)
    
    # Comprueba si hay datos antes de crear la tabla.
    if len(data) > 2 and len(data[0]) > 0:
        table = doc.Tables.Add(table_range, len(data), len(data[0]))
    else:
        print("La tabla markdown en este bot no es reconocida.")
        continue

    # Rellenar la tabla con los datos procesados
    for i, row in enumerate(data):
        for j, cell in enumerate(row):
            table.Cell(i + 1, j + 1).Range.Text = cell

    # Añadir la tabla a la lista de todas las tablas encontradas
    all_tables_data.append(table)

    markdown_link_found_in_table = False  # Añade esta línea para inicializar 'markdown_link_found_in_table' en False

    for i, row_data in enumerate(data):
        for j, cell_data in enumerate(row_data):
            cell = table.Cell(i+1, j+1)
            cell_range = cell.Range

            cell_range.Text = cell_data.strip()
            
            # Manipulación de hipervínculos
            matches = pattern.findall(cell_data)
            
            if matches:
                markdown_link_found_in_table = True  # Se encontró un enlace Markdown en una celda de la tabla

                for text, url in matches:
                    hyperlink_range = cell_range.Duplicate
                  
                    # Limpia el texto de la celda antes de añadir el hipervínculo.
                    cell_range.Text = text.strip()
                    
                    hyperlink_range.Find.Execute(FindText=text)
                    doc.Hyperlinks.Add(Anchor=hyperlink_range, Address=url)
                    
    # Estilo de tabla
    table.Style = "Acc_Table_1"


# Ajusta el tamaño de las celdas a 2.19 cm (convertido a puntos) cuando tiene más de 5 columnas la tabla
cm_to_points = 2.19 * 28.3465  # 1 cm es aproximadamente 28.3465 puntos

# Suponiendo que 'doc' ya está definido y es el documento abierto
# Ajuste del tamaño de las celdas para tablas con 1 a 2 columnas
for table in doc.Tables:
    # Obtener el número de página de la tabla actual
    page_number = table.Range.Information(wdConst.wdActiveEndPageNumber)
    
    # Verificar si "Approval flow" está en la segunda página
    approval_flow_found = False
    for para in doc.Paragraphs:
        if "Approval flow" in para.Range.Text:
            approval_flow_page_number = para.Range.Information(wdConst.wdActiveEndPageNumber)
            if approval_flow_page_number == 2:  # Comprobar si está en la segunda página
                approval_flow_found = True
                break

    # Si se encuentra "Approval flow" en la segunda página, omitir las modificaciones a la tabla
    if approval_flow_found and page_number == 2:
        print(f"Se encontró 'Approval flow' en la segunda página. La tabla en la página {page_number} no será modificada.")
        continue
    
    # Recorrer los encabezados para encontrar "REFERENCES"
    title_found = False
    for para in doc.Paragraphs:
        if "REFERENCES" in para.Range.Text:
            title_page_number = para.Range.Information(wdConst.wdActiveEndPageNumber)
            if title_page_number == page_number:
                title_found = True
                break

    # Si se encuentra el título, omite la modificación de tablas en esa página
    if title_found:
        print(f"Se encontró el título 'REFERENCES TOC HEADING' en la página {page_number}. No se modificará ninguna tabla en esta página.")
        continue
        
    # Tablas con 1 a 2 columnas
    if 1 <= table.Columns.Count <= 2:
        header_height_cm = 2.0  # Ajusta la altura del encabezado en cm
        cell_width_cm = 8.5  # Ajusta el ancho de las celdas en cm

        # Conversión de cm a puntos
        header_height_points = header_height_cm * 28.3465
        cell_width_points = cell_width_cm * 28.3465

        # Ajustar el tamaño de las celdas (tanto el ancho como la altura)
        for i, row in enumerate(table.Rows):
            for cell in row.Cells:
                cell.Width = cell_width_points  # Ancho de todas las celdas
                if i == 0:  # Si es la fila del encabezado
                    cell.Height = header_height_points  # Establece la altura del encabezado
                    cell.HeightRule = wdConst.wdRowHeightExactly  # Establece la altura exacta
                else:
                    cell.HeightRule = wdConst.wdRowHeightAuto  # Altura automática para las otras celdas
        print(f"Tabla de {table.Columns.Count} columnas ajustada con altura de encabezado de {header_height_cm} cm y ancho de celdas de {cell_width_cm} cm.")

    # Tablas con 3 a 5 columnas
    if 3 <= table.Columns.Count <= 5:
        header_height_cm = 2.0  # Ajusta la altura del encabezado en cm
        cell_width_cm = 3.5  # Ajusta el ancho de las celdas en cm

        # Conversión de cm a puntos
        header_height_points = header_height_cm * 28.3465
        cell_width_points = cell_width_cm * 28.3465

        # Ajustar el tamaño de las celdas (tanto el ancho como la altura)
        for i, row in enumerate(table.Rows):
            for cell in row.Cells:
                cell.Width = cell_width_points  # Ancho de todas las celdas
                if i == 0:  # Si es la fila del encabezado
                    cell.Height = header_height_points  # Establece la altura del encabezado
                    cell.HeightRule = wdConst.wdRowHeightExactly  # Establece la altura exacta
                else:
                    cell.HeightRule = wdConst.wdRowHeightAuto  # Altura automática para las otras celdas
        print(f"Tabla de {table.Columns.Count} columnas ajustada con altura de encabezado de {header_height_cm} cm y ancho de celdas de {cell_width_cm} cm.")


    # Tablas con más de 5 columnas
    elif table.Columns.Count > 5:
        # Ajuste manual de celdas para tablas con más de 5 columnas
        cm_to_points = 2.5 * 28.3465  # 1 cm es aproximadamente 28.3465 puntos

        for row in table.Rows:
            for cell in row.Cells:
                try:
                    # Ajusta el ancho de la celda
                    cell.Width = cm_to_points

                    # Verifica el ancho de la celda después del ajuste
                    actual_width = cell.Width

                    # Comprueba si el ancho actual es aproximadamente el esperado
                    if abs(actual_width - cm_to_points) < 1:
                        print("La celda se ha ajustado correctamente.")
                    else:
                        print("La celda no se ajustó al tamaño esperado.")

                except Exception as e:
                    print(f"Se produjo un error al intentar ajustar el tamaño de la celda: {e}")
        print(f"Tabla de {table.Columns.Count} columnas ajustada manualmente con un ancho de celdas de 2.4 cm.")

# Inicializa 'sheet_resized' como False



print("Aumentando el tamaño de la hoja microsoft word en el caso de pueda ser necesario")

# Recorrer las tablas y aumentar el tamaño de la hoja si hay más de 5 columnas
for table in doc.Tables:
    if table.Columns.Count > 5:
        try:
            # Aumentar el tamaño de la hoja en 2.54 cm (equivalente a 1 pulgada)
            points_in_cm = 5.45 * 28.3465  # 1 cm = 28.3465 puntos
            doc.PageSetup.PageWidth = doc.PageSetup.PageWidth + points_in_cm
            doc.PageSetup.PageHeight = doc.PageSetup.PageHeight + points_in_cm
            print(f"The size of the sheet has been increased. New width: {doc.PageSetup.PageWidth}, Nuevo alto: {doc.PageSetup.PageHeight}")
            sheet_resized = True
            break  # Salir del bucle después de la primera tabla encontrada
        except Exception as e:
            print(f"An error occurred while trying to adjust the sheet size: {e}")

if sheet_resized:
    print("The document sheet size has been resized.")
else:
    print("The size of the document sheet has not been altered.")



print("Sección comprobación tamaño de hoja y cambio de shape de texto") 
# Definir factor de conversión de cm a puntos (Word usa puntos como medida)
cm_to_points = 28.3465

# Tamaño A3 en cm
a3_width_in_cm = 29.7  # Ancho en cm
a3_height_in_cm = 42.0  # Alto en cm

# Convertir tamaño A3 a puntos
a3_width_in_points = a3_width_in_cm * cm_to_points
a3_height_in_points = a3_height_in_cm * cm_to_points

# Tamaño de ancho absoluto para el cuadro de texto en cm (26.65 cm)
new_width_in_cm = 26.65
new_width_in_points = new_width_in_cm * cm_to_points

# Variable para verificar si el tamaño de la hoja fue ajustado
sheet_resized = False

# Recorrer las tablas y cambiar el tamaño de la hoja a A3 si hay más de 5 columnas
for table in doc.Tables:
    if table.Columns.Count > 5:
        try:
            # Cambiar el tamaño de la hoja a A3
            doc.PageSetup.PageWidth = a3_width_in_points
            doc.PageSetup.PageHeight = a3_height_in_points
            print(f"Tamaño de las hojas cambiado a A3: {a3_width_in_cm} cm x {a3_height_in_cm} cm.")
            sheet_resized = True
            break  # Salir del bucle después de ajustar la primera hoja con una tabla grande
        except Exception as e:
            print(f"Error al intentar ajustar el tamaño de las hojas: {e}")

# Si el tamaño de la hoja fue ajustado, buscar el cuadro de texto en la última hoja
if sheet_resized:
    print("Ajustando el cuadro de texto en la última hoja...")

    # Obtener todas las formas (Shapes) en el documento completo
    shapes = doc.Shapes

    if shapes.Count > 0:
        # Iterar por todas las Shapes para encontrar el cuadro de texto en la última sección
        for shape in shapes:
            # Verificar si el Shape es un cuadro de texto y si está en la última sección
            if shape.TextFrame.HasText:
                shape_range = shape.Anchor
                if shape_range.Sections(1).Index == doc.Sections.Count:
                    try:
                        # Ajustar el ancho del cuadro de texto
                        shape.LockAspectRatio = False  # Desbloquear la relación de aspecto
                        shape.Width = new_width_in_points  # Establecer el ancho en puntos
                        print(f"El ancho del cuadro de texto ha sido ajustado a {new_width_in_cm} cm.")
                        break  # Salir del bucle una vez que el cuadro de texto en la última sección sea encontrado
                    except Exception as e:
                        print(f"Error al ajustar el tamaño del cuadro de texto: {e}")
                else:
                    print("No se encontraron cuadros de texto en la última hoja.")
    else:
        print("No se encontraron cuadros de texto en el documento.")
else:
    print("El tamaño de las hojas no fue alterado, no se ajustará ningún cuadro de texto.")



print("medidas personalizadas para el cuadro contenedor de las formas shape en páginas A3")
active_document = word_app.ActiveDocument

# Factor de conversión de puntos a centímetros
points_to_cm = 0.0352778
cm_to_points = 28.3464567  # Conversión de cm a puntos

# Verificar el tamaño de la página en centímetros
page_width_cm = active_document.PageSetup.PageWidth * points_to_cm
page_height_cm = active_document.PageSetup.PageHeight * points_to_cm

print(f"Tamaño de la página: {page_width_cm:.2f} cm x {page_height_cm:.2f} cm")

# Verificar si las dimensiones coinciden con el tamaño A3 (29.7 cm x 42.0 cm)
if abs(page_width_cm - 29.7) < 0.1 and abs(page_height_cm - 42.0) < 0.1:
    print("La página está en tamaño A3, aplicando configuración...")
    
    # Establecer configuración de Layout Position y Layout Size solo para el cuadro contenedor
    if shapes.Count > 0:
        container_shape = None

        for shape in shapes:
            # Verificar si la forma tiene imágenes anidadas
            if shape.GroupItems.Count > 1:  # Verificar si es un grupo con más de 1 elemento (imágenes anidadas)
                print(f"Cuadro contenedor encontrado: {shape.Name}, configurando tamaño y posición.")
                container_shape = shape
                break  # Detener el bucle después de encontrar el cuadro contenedor

        if container_shape:
            try:
                # Establecer Layout Position
                container_shape.Left = 0  # Posición horizontal absoluta
                
                # Convertir 3.93 cm a puntos y verificar el valor
                top_cm = 3.93
                top_points = top_cm * cm_to_points
                print(f"Posición Top (vertical) en centímetros: {top_cm} cm")
                print(f"Posición Top (vertical) en puntos: {top_points}")

                # Establecer la posición relativa de la página
                container_shape.RelativeVerticalPosition = constants.wdRelativeVerticalPositionPage
                container_shape.Top = top_points  # Posición vertical absoluta en puntos

                # Imprimir la posición en centímetros después de establecerla en puntos
                print(f"Posición Left (horizontal): {container_shape.Left * points_to_cm:.2f} cm")
                print(f"Posición Top (vertical): {container_shape.Top * points_to_cm:.2f} cm")

                # Verificar la propiedad RelativeVerticalPosition
                print(f"RelativeVerticalPosition: {container_shape.RelativeVerticalPosition}")

                # Establecer Layout Size
                container_shape.LockAspectRatio = False  # Desbloquear la relación de aspecto
                container_shape.Height = 28.16 * cm_to_points  # Establecer la altura en puntos
                container_shape.Width = 25.2 * cm_to_points  # Establecer el ancho en puntos

                # Imprimir el tamaño en centímetros
                print(f"Altura (Height): {container_shape.Height * points_to_cm:.2f} cm")
                print(f"Ancho (Width): {container_shape.Width * points_to_cm:.2f} cm")

                print("Configuración de Layout Position y Layout Size establecida para el cuadro contenedor.")
            except Exception as e:
                print(f"Error al establecer la configuración de Layout Position y Layout Size: {e}")
        else:
            print("No se encontró ningún cuadro contenedor para aplicar la configuración.")
else:
    print("La página no está en tamaño A3. No se aplicarán cambios.")


print("Verificando tamaño de la portada y listando shapes...")

def verificar_portada_y_shapes(doc):
    # Factor de conversión de puntos a cm
    points_to_cm = 0.0352778
    cm_to_points = 28.3464567

    try:
        # Verificar tamaño de la página
        page_width_cm = doc.PageSetup.PageWidth * points_to_cm
        page_height_cm = doc.PageSetup.PageHeight * points_to_cm

        print(f"Tamaño de la página: {page_width_cm:.2f} cm x {page_height_cm:.2f} cm")
        if abs(page_width_cm - 29.7) < 0.1 and abs(page_height_cm - 42.0) < 0.1:
            print("La página está en tamaño A3.")
        else:
            print("La página no está en tamaño A3.")
        
        # Iterar sobre shapes y obtener información detallada
        for shape in shapes:
            if shape.Type == 6:  # Si es un grupo
                for item in shape.GroupItems:
                    if item.Type == 17:  # wdInlineShapeTextBox
                        pass  # Aquí puedes agregar el código para manejar el TextBox
                    elif item.Type == 3:  # wdInlineShapePicture
                        pass  # Aquí puedes agregar el código para manejar la imagen
            elif shape.Name == "Picture 14":  # Ajustar la posición de la imagen específica
                target_position_cm = 20
                current_position_cm = shape.Left * points_to_cm
                distance_to_move_cm = target_position_cm - current_position_cm
                distance_to_move_points = distance_to_move_cm * cm_to_points
                shape.IncrementLeft(distance_to_move_points)
                print(f"Imagen 'Picture 14' movida a {target_position_cm} cm desde la izquierda.")
                
            # Ajustar posición de Rectangle: Rounded Corners 4
            if shape.Name == "Rectangle: Rounded Corners 4":
                target_position_cm = 20
                current_position_cm = shape.Left * points_to_cm
                distance_to_move_cm = target_position_cm - current_position_cm
                distance_to_move_points = distance_to_move_cm * cm_to_points
                shape.IncrementLeft(distance_to_move_points)
                print(f"'{shape.Name}' movido a {target_position_cm} cm desde la izquierda.")    
    except Exception as e:
        print(f"Error en verificación de la portada o shapes: {e}")

# Ejecutar función de verificación
verificar_portada_y_shapes(doc)
        

print("Configuración personalizada de tabla en el encabezado para páginas A3")

# Factor de conversión de puntos a centímetros
points_to_cm = 0.0352778
cm_to_points = 28.3464567  # Conversión de cm a puntos

# Verificar el tamaño de la página en centímetros
page_width_cm = active_document.PageSetup.PageWidth * points_to_cm
page_height_cm = active_document.PageSetup.PageHeight * points_to_cm

print(f"Tamaño de la página: {page_width_cm:.2f} cm x {page_height_cm:.2f} cm")

# Verificar si las dimensiones coinciden con el tamaño A3 (29.7 cm x 42.0 cm)
if abs(page_width_cm - 29.7) < 0.1 and abs(page_height_cm - 42.0) < 0.1:
    print("La página está en tamaño A3, aplicando configuración de ancho de tabla...")

    # Intentar acceder a los encabezados de la sección 2 (Principal y Primera Página)
    try:
        # Encabezado Principal
        header_primary = active_document.Sections(2).Headers(win32.constants.wdHeaderFooterPrimary)
        tables_in_primary_header = [tbl for tbl in header_primary.Range.Tables]

        # Encabezado de la Primera Página
        header_first_page = active_document.Sections(2).Headers(win32.constants.wdHeaderFooterFirstPage)
        tables_in_first_page_header = [tbl for tbl in header_first_page.Range.Tables]

        # Aplicar ancho de 24 cm en el encabezado principal
        if tables_in_primary_header:
            for table in tables_in_primary_header:
                table.PreferredWidth = 24 * cm_to_points
                print("Ancho preferido de la tabla en el encabezado principal ajustado a 24 cm.")
        else:
            print("No se encontraron tablas en el encabezado principal de la sección 2.")

        # Aplicar ancho de 24 cm en el encabezado de la primera página
        if tables_in_first_page_header:
            for table in tables_in_first_page_header:
                table.PreferredWidth = 24 * cm_to_points
                print("Ancho preferido de la tabla en el encabezado de la primera página del documento word ajustado a 24 cm.")
        else:
            print("No se encontraron tablas en el encabezado de la primera página de la sección 2.")

    except Exception as e:
        print(f"Error al acceder al encabezado o ajustar el ancho de la tabla: {e}")
else:
    print("La página no está en tamaño A3. No se aplicarán cambios en los encabezados del documento Word.")



print("Añadiendo menciones azure wiki de color rojo")
def resaltar_menciones(doc):
    ###Busca menciones y las cambia a color rojo, eliminando paréntesis si los hay
    print("Añadiendo menciones azure de color rojo")

    # Expresión regular para identificar las menciones
    # Casos a cubrir:
    # - @(ID) donde ID es una secuencia de caracteres
    # - @Nombre
    mention_pattern = re.compile(r'@(?:\(([^\)]+)\)|(\w+))')

    # Recorre cada párrafo en el documento
    for paragraph in doc.Paragraphs:
        paragraph_text = paragraph.Range.Text

        # Buscar todas las menciones en el párrafo
        matches = list(mention_pattern.finditer(paragraph_text))

        # Inicializar un offset para ajustar las posiciones después de reemplazos
        offset = 0

        for match in matches:
            # Obtener las posiciones inicial y final ajustadas
            start = match.start() + offset
            end = match.end() + offset

            full_mention = match.group(0)
            mention_content = match.group(1) if match.group(1) else match.group(2)

            # Crear la nueva mención sin paréntesis
            new_mention = f"@{mention_content}"

            # Calcular la diferencia de longitud para ajustar el offset
            length_diff = len(new_mention) - len(full_mention)

            # Crear un Range para la mención
            mention_range = paragraph.Range.Duplicate
            mention_range.SetRange(paragraph.Range.Start + start, paragraph.Range.Start + end)

            # Reemplazar el texto en el rango de la mención
            mention_range.Text = new_mention

            # Aplicar el color rojo a la mención
            mention_range.Font.Color = 255  # wdColorRed in Word

            # Actualizar el offset
            offset += length_diff
            
# Llama a la función para resaltar menciones
resaltar_menciones(doc)
            


print("reemplazo por salto de párrafos con (br)")
def reemplazar_br_por_salto(doc):
    ###Reemplaza todas las ocurrencias de (br) por un salto de párrafo en el documento.
    print("Reemplazando '(br)' por saltos de párrafo")

    # Recorre todos los párrafos en el documento
    for paragraph in doc.Paragraphs:
        paragraph_text = paragraph.Range.Text

        # Verifica si el texto contiene '(br)'
        if '(br)' in paragraph_text:
            # Reemplaza '(br)' por un salto de párrafo
            paragraph.Range.Text = paragraph_text.replace('(br)', '\\r')  #   r es un salto de párrafo en Word

# Llama a la función para reemplazar '(br)' por saltos de párrafo
reemplazar_br_por_salto(doc)


print("Formatting blue label b color blue /b")
# Patrón para buscar las etiquetas con diferentes colores
pattern = re.compile(r'\(b\)\(a color:blue\)(.*?)\(/b\)')

## Recorrer todos los párrafos del documento
for paragraph in doc.Paragraphs:
    paragraph_text = paragraph.Range.Text
    # Buscar todas las coincidencias del patrón en el párrafo
    matches = pattern.finditer(paragraph_text)
    
    for match in matches:
        matched_text = match.group(1)
        
        # Encontrar la posición del texto en el párrafo
        start = match.start(1)
        end = match.end(1)
        
        # Crear un rango para el texto
        text_range = doc.Range(paragraph.Range.Start + start, paragraph.Range.Start + end)
        
        # Aplicar el formato
        text_range.Font.Bold = True
        text_range.Font.Color = win32.constants.wdColorBlue

        # Ajustar índices para eliminar las etiquetas
        start_tag_range = doc.Range(paragraph.Range.Start + match.start(), paragraph.Range.Start + start)
        end_tag_range = doc.Range(paragraph.Range.Start + end, paragraph.Range.Start + match.end())

        # Eliminar las etiquetas
        start_tag_range.Delete()
        end_tag_range.Delete()


print("strike format word")

for para in doc.Paragraphs:
    original_text = para.Range.Text
    matches = list(re.finditer(r'~~(.*?)~~', original_text))

    # Procesar cada coincidencia en reversa para no desajustar las posiciones subsiguientes
    for match in reversed(matches):
        # Obtener las posiciones inicial y final del texto a tachar (incluyendo los símbolos ~~)
        start = match.start()
        end = match.end()

        # Obtener el texto a tachar (sin los símbolos ~~)
        text_to_strike = match.group(1)

        # Crear un rango para el texto completo a modificar (incluyendo los símbolos ~~)
        full_range = doc.Range(para.Range.Start + start, para.Range.Start + end)
        
        # Primero, aplicar el formato tachado al texto
        full_range.Font.StrikeThrough = True

        # Luego, reemplazar el texto completo (incluyendo los símbolos ~~) solo por el texto a tachar
        full_range.Text = text_to_strike

        # Actualizar el texto original para reflejar el cambio
        original_text = original_text[:start] + text_to_strike + original_text[end:]



print("Applying blockquote format in the word document")

# Permite uno o más paréntesis de cierre seguidos de espacio
blockquote_pattern = r'^[\)>]+\s*'

for para in doc.Paragraphs:
    full_text = para.Range.Text
    # Comprobamos si arranca con uno o más ")" y un espacio
    match = re.match(blockquote_pattern, full_text)
    if match:
        # Elimina los paréntesis y el espacio inicial
        delete_range = para.Range.Duplicate
        delete_range.End = delete_range.Start + match.end()
        delete_range.Delete()

        # Aseguramos el estilo "Normal" antes de aplicar bordes
        para.Range.Style = word_app.ActiveDocument.Styles("Normal")

        # Aplicamos el borde izquierdo
        border = para.Range.ParagraphFormat.Borders(constants.wdBorderLeft)
        border.LineStyle = constants.wdLineStyleSingle
        border.LineWidth = constants.wdLineWidth225pt  # borde grueso
        border.Color = win32api.RGB(234, 234, 234)
        
        # Aplicamos la tabulación para diferenciar el texto
        indent_level = 20 * match.end()  # 20 puntos por cada paréntesis
        para.Range.ParagraphFormat.LeftIndent = indent_level  # Aplica sangría


        # Log
        print(f"Formato aplicado al párrafo: {full_text.strip()}")



print("gestiona etiquetas html sup")
# Patrón para pares completos: captura lo que hay entre (sup) y (/sup)
pair_re = re.compile(r'\(sup\)(.*?)\(/sup\)', re.DOTALL)
# Patrón para aperturas sueltas
open_re = re.compile(r'\(sup\)')

for para in doc.Paragraphs:
    rng = para.Range
    text = rng.Text  # incluye el carácter de párrafo al final

    # 1) Procesar todos los pares explícitos (sup)…(/sup)
    for m in reversed(list(pair_re.finditer(text))):
        full_start, full_end = m.span(0)
        inner_text = m.group(1)

        # Rango absoluto en el documento
        sup_start = rng.Start + full_start + len("(sup)")
        sup_end   = sup_start + len(inner_text)

        # Aplica superíndice al contenido
        doc.Range(sup_start, sup_end).Font.Superscript = True

        # Borra primero la etiqueta de cierre
        doc.Range(rng.Start + full_start + len("(sup)") + len(inner_text),
                  rng.Start + full_end).Delete()
        # Luego borra la etiqueta de apertura
        doc.Range(rng.Start + full_start,
                  rng.Start + full_start + len("(sup)")).Delete()

        # Refresca texto del párrafo tras las eliminaciones
        text = rng.Text

    # 2) Procesar aperturas sueltas: (sup)…<fin de párrafo>
    for m in reversed(list(open_re.finditer(text))):
        start = m.start()
        next_idx = start + len("(sup)")

        # Si tras (sup) hay un espacio, solo borramos la etiqueta
        if next_idx < len(text) and text[next_idx].isspace():
            doc.Range(rng.Start + start,
                      rng.Start + start + len("(sup)")).Delete()
            text = rng.Text
            continue

        # Si no hay espacio, aplicamos superíndice hasta fin de párrafo
        sup_start = rng.Start + start + len("(sup)")
        sup_end   = rng.End - 1  # antes del salto de párrafo

        doc.Range(sup_start, sup_end).Font.Superscript = True

        # Borra la etiqueta de apertura
        doc.Range(rng.Start + start,
                  rng.Start + start + len("(sup)")).Delete()

        # Refresca texto del párrafo tras las eliminaciones
        text = rng.Text

print("✅ Formateo de superíndices con (sup)…(/sup) y aperturas huérfanas completado.")



print("gestiona etiquetas html sub")
# Patrón para pares completos: captura lo que hay entre (sub) y (/sub)
pair_re = re.compile(r'\(sub\)(.*?)\(/sub\)', re.DOTALL)
# Patrón para aperturas sueltas
open_re = re.compile(r'\(sub\)')

for para in doc.Paragraphs:
    rng = para.Range
    text = rng.Text  # incluye el carácter de párrafo al final

    # 1) Procesar todos los pares explícitos (sub)…(/sub)
    for m in reversed(list(pair_re.finditer(text))):
        full_start, full_end = m.span(0)
        inner_text = m.group(1)

        # Rango absoluto en el documento
        sub_start = rng.Start + full_start + len("(sub)")
        sub_end   = sub_start + len(inner_text)

        # Aplica subíndice al contenido y aumenta el tamaño
        sub_range = doc.Range(sub_start, sub_end)
        sub_range.Font.Subscript = True
        sub_range.Font.Size = 12  # Aumenta el tamaño (ajusta según necesidad)

        # Borra primero la etiqueta de cierre
        doc.Range(rng.Start + full_end - len("(/sub)"), rng.Start + full_end).Delete()
        # Luego borra la etiqueta de apertura
        doc.Range(rng.Start + full_start, rng.Start + full_start + len("(sub)")).Delete()

        # Refresca texto del párrafo tras las eliminaciones
        text = rng.Text

    # 2) Procesar aperturas sueltas: (sub)…<fin de párrafo>
    for m in reversed(list(open_re.finditer(text))):
        start = m.start()
        next_idx = start + len("(sub)")

        # Si tras (sub) hay un espacio (o más), solo borramos la etiqueta
        if next_idx < len(text) and text[next_idx].isspace():
            doc.Range(rng.Start + start, rng.Start + start + len("(sub)")).Delete()
            text = rng.Text
            continue

        # Si no hay espacio, verificamos que haya texto válido (como H2O) pegado a (sub)
        if next_idx < len(text):
            sub_start = rng.Start + start + len("(sub)")
            sub_end   = rng.End - 1  # antes del salto de párrafo

            # Aplica subíndice y aumenta el tamaño
            sub_range = doc.Range(sub_start, sub_end)
            sub_range.Font.Subscript = True
            sub_range.Font.Size = 15  # Aumenta el tamaño (ajusta según necesidad)

        # Borra la etiqueta de apertura
        doc.Range(rng.Start + start, rng.Start + start + len("(sub)")).Delete()

        # Refresca texto del párrafo tras las eliminaciones
        text = rng.Text

print("✅ Formateo de subíndices con (sub)…(/sub) y aperturas huérfanas completado.")


print("Rename ficheros formato imagen en mayuscula lo convierte en minusculas")

# Función para renombrar archivos en un directorio
def rename_files_in_directory(directory):
    for filename in os.listdir(directory):
        # Dividir el nombre del archivo y la extensión
        file_root, file_extension = os.path.splitext(filename)
        # Si la extensión del archivo es .PNG o .JPG
        if file_extension.upper() in ['.PNG', '.JPG']:
            # Reemplazar espacios con guiones en el nombre del archivo
            new_file_root = file_root.replace(' ', '_')
            new_filename = f"{new_file_root}{file_extension.lower()}"

            # Renombrar el archivo
            try:
                os.rename(os.path.join(directory, filename), os.path.join(directory, new_filename))
                print(f"¡File successfully renamed!")
            except Exception as e:
                print(f"Error renaming {filename}: {e}")

# Obtener la ruta del directorio actual y unirla con .attachments
current_directory = os.path.dirname(os.path.abspath(__file__))
directory_path = os.path.join(current_directory, '.attachments')

# Llamar a la función con la ruta del directorio
rename_files_in_directory(directory_path)



print("Descarga imágenes y luego las pone en el documento Word, manejando enlaces con y sin PIPES.")

image_pattern_with_pipes = re.compile(r'\|!\[([^\]]*)\]\(([^)]+)\)\|')
image_pattern_without_pipes = re.compile(r'!\[([^\]]*)\]\(([^)]+)\)')

script_dir = os.path.dirname(os.path.abspath(__file__))
attachments_dir = os.path.join(script_dir, ".attachments")

if not os.path.isdir(attachments_dir):
    print(f"El directorio de imágenes no existe: {attachments_dir}")
else:
    print(f"Directorio .attachments encontrado: {attachments_dir}")
    attachment_files = os.listdir(attachments_dir)

def insert_image_with_pipes(doc, match_range, image_path):

    try:
        # Eliminar el texto de la coincidencia encontrada
        match_range.Delete()
        
        # Crear un rango de inserción
        insertion_range = match_range.Duplicate
        insertion_range.Collapse(0)
        
        # Máxima altura permitida (6 cm)
        max_height = 6 * 28.3465  # 6 cm en puntos
        max_width_cm = 15  # Umbral de ancho en cm
        max_width_points = max_width_cm * 28.3465  # Convertir a puntos

        # Crear tabla temporal para medir la imagen
        temp_table = doc.Tables.Add(insertion_range, 1, 1)
        temp_table.Borders.Enable = False  # Sin bordes temporales
        
        # Insertar imagen temporal para medir dimensiones
        image_cell = temp_table.Cell(1, 1)
        image = image_cell.Range.InlineShapes.AddPicture(
            FileName=image_path,
            LinkToFile=False,
            SaveWithDocument=True
        )
        
        # Ajustar solo la altura de la imagen a 6 cm (sin tocar el ancho)
        if image.Height > max_height:
            image.Height = max_height  # Solo altura ajustada
        
        # Guardar ancho ajustado
        image_width = image.Width  # Ancho permanece intacto
        
        # Eliminar tabla temporal
        temp_table.Delete()
        
        # Crear nueva tabla según el ancho ajustado
        if image_width < max_width_points:  # Imagen más estrecha que 15 cm
            # Crear tabla con 3 filas y 1 columna (fila inferior dividida en 2 columnas)
            table = doc.Tables.Add(insertion_range, 3, 1)
            table.Borders.Enable = False  # Mostrar bordes para diagnóstico
            
            # Configurar fila superior vacía (pipeline superior omitido para < 15 cm)
            table.Rows(1).Height = 1
            table.Cell(1, 1).Range.Text = ""  # Fila superior sin contenido
            
            # Fila central para la imagen
            image_cell = table.Cell(2, 1)
            image = image_cell.Range.InlineShapes.AddPicture(
                FileName=image_path,
                LinkToFile=False,
                SaveWithDocument=True
            )
            
            # Ajustar solo la altura (sin modificar el ancho)
            if image.Height > max_height:
                image.Height = max_height  # Solo altura ajustada
            
            # Fila inferior con dos pipelines
            table.Cell(3, 1).Split(1, 2)  # Dividir celda en dos columnas
            
            # Ajustar el ancho de la tabla al ancho de la imagen
            table.PreferredWidth = image_width  # Ancho total de la tabla igual al ancho de la imagen
            
            # Ajustar las celdas izquierda y derecha
            # Hacemos que la celda derecha sea proporcional al contenido del pipeline
            table.Cell(3, 1).Width = image_width / 2 - 1  # Reducir ancho de la celda izquierda
            table.Cell(3, 2).Width = image_width / 2 - 1  # Reducir ancho de la celda derecha
            
            # Ajustar márgenes internos de las celdas
            table.Cell(3, 1).Range.ParagraphFormat.LeftIndent = 0
            table.Cell(3, 1).Range.ParagraphFormat.RightIndent = 0
            
            table.Cell(3, 2).Range.ParagraphFormat.LeftIndent = 0  # Sin margen interno izquierdo
            table.Cell(3, 2).Range.ParagraphFormat.RightIndent = 0
            
            # Pipeline inferior izquierdo
            table.Cell(3, 1).Range.Text = "|"  
            table.Cell(3, 1).Range.ParagraphFormat.Alignment = wdConst.wdAlignParagraphLeft
            
            # Pipeline inferior derecho
            table.Cell(3, 2).Range.Text = "|"
            table.Cell(3, 2).Range.ParagraphFormat.Alignment = wdConst.wdAlignParagraphLeft  # Pegado al borde izquierdo
        else:
            # Imagen más ancha o igual a 15 cm: tabla simple con pipelines en esquinas superior e inferior izquierdas
            table = doc.Tables.Add(insertion_range, 3, 1)
            table.Borders.Enable = False  # Mostrar bordes para diagnóstico
            
            # Fila superior con pipeline
            table.Cell(1, 1).Range.Text = "|"  # Pipeline superior izquierdo
            table.Cell(1, 1).Range.ParagraphFormat.Alignment = wdConst.wdAlignParagraphLeft
            table.Cell(1, 1).Range.ParagraphFormat.SpaceBefore = 0
            table.Cell(1, 1).Range.ParagraphFormat.SpaceAfter = 0
            
            # Fila central para la imagen
            image_cell = table.Cell(2, 1)
            image = image_cell.Range.InlineShapes.AddPicture(
                FileName=image_path,
                LinkToFile=False,
                SaveWithDocument=True
            )
            
            # Ajustar solo la altura (sin modificar el ancho)
            if image.Height > max_height:
                image.Height = max_height  # Solo altura ajustada
            
            # Fila inferior con pipeline
            table.Cell(3, 1).Range.Text = "|"  # Pipeline inferior izquierdo
            table.Cell(3, 1).Range.ParagraphFormat.Alignment = wdConst.wdAlignParagraphLeft
            table.Cell(3, 1).Range.ParagraphFormat.SpaceBefore = 0
            table.Cell(3, 1).Range.ParagraphFormat.SpaceAfter = 0

        # Ajustar automáticamente el ancho de las columnas
        table.Columns.AutoFit()
        return True
    except Exception as e:
        print(f"Error al insertar la imagen: {e}")
        return False

def insert_image_without_pipes(doc, match_range, image_path):
    try:
        match_range.Delete()
        image = match_range.InlineShapes.AddPicture(
            FileName=image_path,
            LinkToFile=False,
            SaveWithDocument=True
        )
        max_height = 6 * 28.3465
        if image.Height > max_height:
            image.Height = max_height
        match_range.ParagraphFormat.Alignment = win32.constants.wdAlignParagraphCenter
        return True
    except Exception as e:
        print(f"Error al insertar la imagen: {e}")
        return False

try:
    paragraphs = list(doc.Paragraphs)
    for paragraph in paragraphs:
        match_with_pipes = image_pattern_with_pipes.search(paragraph.Range.Text)
        match_without_pipes = image_pattern_without_pipes.search(paragraph.Range.Text)

        if match_with_pipes or match_without_pipes:
            if match_with_pipes:
                description = match_with_pipes.group(1)
                image_path_markdown = match_with_pipes.group(2).lstrip('/')
                has_pipes = True
                match = match_with_pipes
            else:
                description = match_without_pipes.group(1)
                image_path_markdown = match_without_pipes.group(2).lstrip('/')
                has_pipes = False
                match = match_without_pipes

            image_path_markdown = re.sub(r'\s*=\d+[xX]\d*', '', image_path_markdown)
            if image_path_markdown.startswith(("http", "https")):
                continue

            image_path_markdown = urllib.parse.unquote(image_path_markdown)
            file_name_ext = os.path.basename(image_path_markdown)
            file_name_ext = file_name_ext.replace(' ', '_')  # Ensure spaces are replaced with underscores

            if file_name_ext in attachment_files:
                image_path = os.path.join(attachments_dir, file_name_ext)
                match_start = match.start()
                match_end = match.end()
                match_range = paragraph.Range.Duplicate
                match_range.Start = paragraph.Range.Start + match_start
                match_range.End = paragraph.Range.Start + match_end

                if has_pipes:
                    if insert_image_with_pipes(doc, match_range, image_path):
                        print(f"Imagen insertada correctamente con pipes: {file_name_ext}")
                    else:
                        print(f"Imagen saltada o error al insertar: {file_name_ext}")
                else:
                    if insert_image_without_pipes(doc, match_range, image_path):
                        print(f"Imagen insertada correctamente sin pipes: {file_name_ext}")
                    else:
                        print(f"Imagen saltada o error al insertar: {file_name_ext}")

except Exception as e:
    print(f"Se capturó una excepción al procesar imágenes: {e}")
    try:
        print(f"Excepción capturada en el párrafo: {paragraph.Range.Text}")
    except:
        print("No se pudo obtener el texto del párrafo.")
    print(f"Detalles de la excepción: {e}")

   
   

print("imagenes con enlaces html página imgbb img SRC (gestionando casos con y sin pipes)")

def verificar_conexion_ibb():
    try:
        response = requests.get("https://ibb.co/", timeout=10, headers={'User-Agent': 'Mozilla/5.0'})
        if response.status_code == 200:
            print("✅ Conexión a ibb.co verificada correctamente")
            return True
        else:
            print(f"⚠️ ibb.co respondió con código {response.status_code}, continuando sin verificación")
            return False
    except requests.exceptions.RequestException as e:
        print(f"⚠️ Error al conectar con ibb.co: {e}. Continuando sin verificación")
        return False

def procesar_imagenes_html(doc):
    # Verificar conexión a ibb.co al inicio
    ibb_disponible = verificar_conexion_ibb()
    
    # Intentar importar PIL para obtener dimensiones
    try:
        from PIL import Image
        pillow_installed = True
    except ImportError:
        pillow_installed = False

    # Directorio .attachments
    script_dir = os.path.dirname(os.path.abspath(__file__))
    attachments_dir = os.path.join(script_dir, ".attachments")
    os.makedirs(attachments_dir, exist_ok=True)
    print(f"📂 Directorio .attachments: {attachments_dir}")

    # Patrones: con y sin pipes
    pipe_pattern = re.compile(r"\|\(a href=\\"[^\\"]+\\"\)\(img src=\\"([^\\"]+)\\"[^>]*?/\\)\\(/a\)\\|", re.IGNORECASE)
    html_pattern = re.compile(r"\(a href=\\"[^\\"]+\\"\\)\(img src=\\"([^\\"]+)\\"[^>]*?/\\)\\(/a\)", re.IGNORECASE)

    # Constantes de tamaño y pipelines
    MAX_HEIGHT = 6 * 28.3465
    PIPE_FONT = 12
    PIPE_ROW_H = PIPE_FONT + 2

    paragraphs = list(doc.Paragraphs)
    session = requests.Session()
    headers = {'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'}
    idx = 0

    def insert_with_pipes(range_obj, path, width, height):
        try:
            range_obj.Delete()
            range_obj.Collapse(win32.constants.wdCollapseStart)
            # Evaluar tamaño
            max_h = MAX_HEIGHT
            max_w = 15 * 28.3465

            # Asegurar que width y height sean valores válidos
            if width <= 0 or height <= 0:
                width, height = 1000, 800  # Valores por defecto si las dimensiones fallan

            # Decidir diseño según proporción: vertical (pipes a los lados) o horizontal (pipes arriba/abajo)
            if height > width:
                # Imagen más alta que ancha: pipes a los lados (tabla 1x3)
                tbl = doc.Tables.Add(range_obj, 1, 3)
                tbl.Borders.Enable = False
                tbl.Rows.AllowBreakAcrossPages = False
                # Izquierda
                c1 = tbl.Cell(1, 1)
                c1.Range.Text = "|"
                c1.Range.Font.Size = PIPE_FONT
                c1.Range.ParagraphFormat.Alignment = win32.constants.wdAlignParagraphCenter
                # Imagen
                c2 = tbl.Cell(1, 2)
                pic = c2.Range.InlineShapes.AddPicture(path, False, True)
                if pic.Height > max_h:
                    pic.Height = max_h
                c2.Range.ParagraphFormat.Alignment = win32.constants.wdAlignParagraphCenter
                # Derecha
                c3 = tbl.Cell(1, 3)
                c3.Range.Text = "|"
                c3.Range.Font.Size = PIPE_FONT
                c3.Range.ParagraphFormat.Alignment = win32.constants.wdAlignParagraphCenter
                tbl.Columns.AutoFit()
            else:
                # Imagen más ancha que alta: pipes arriba y abajo (tabla 3x1)
                tbl = doc.Tables.Add(range_obj, 3, 1)
                tbl.Borders.Enable = False
                tbl.Rows.AllowBreakAcrossPages = False
                # Arriba
                top = tbl.Cell(1, 1)
                top.Range.Text = "|"
                top.Range.Font.Size = PIPE_FONT
                top.Range.ParagraphFormat.Alignment = win32.constants.wdAlignParagraphLeft
                top.VerticalAlignment = win32.constants.wdCellAlignVerticalBottom
                top.Range.ParagraphFormat.SpaceAfter = 0
                top.Range.ParagraphFormat.SpaceBefore = 0
                tbl.Rows(1).HeightRule = win32.constants.wdRowHeightExactly
                tbl.Rows(1).Height = PIPE_FONT - 2  # Reducido para acercar el pipe
                # Medio
                mid = tbl.Cell(2, 1)
                pic = mid.Range.InlineShapes.AddPicture(path, False, True)
                if pic.Height > max_h:
                    pic.Height = max_h
                mid.Range.ParagraphFormat.Alignment = win32.constants.wdAlignParagraphCenter
                # Abajo
                bot = tbl.Cell(3, 1)
                bot.Range.Text = "|"
                bot.Range.Font.Size = PIPE_FONT
                bot.Range.ParagraphFormat.Alignment = win32.constants.wdAlignParagraphLeft
                tbl.Rows(3).HeightRule = win32.constants.wdRowHeightExactly
                tbl.Rows(3).Height = PIPE_ROW_H
                tbl.Columns.AutoFit()
            return True
        except Exception as e:
            print(f"Error insert_with_pipes: {e}")
            return False

    def insert_without_pipes(range_obj, path):
        try:
            range_obj.Delete()
            range_obj.Collapse(win32.constants.wdCollapseStart)
            pic = range_obj.InlineShapes.AddPicture(path, False, True)
            if pic.Height > MAX_HEIGHT:
                pic.Height = MAX_HEIGHT
            range_obj.ParagraphFormat.Alignment = win32.constants.wdAlignParagraphCenter
            return True
        except Exception as e:
            print(f"Error insert_without_pipes: {e}")
            return False

    def descargar_imagen(img_url, session, headers):
        try:
            # Si la URL es de ibb.co y no pudimos verificar la conexión, intentar de todas formas
            if "ibb.co" in img_url.lower() and not ibb_disponible:
                print(f"⚠️ Intentando descargar de ibb.co sin verificación previa: {img_url}")
            
            r = session.get(img_url, headers=headers, stream=True, timeout=15)
            if r.status_code != 200:
                print(f"❌ Error HTTP {r.status_code} al descargar: {img_url}")
                return None, None, None
                
            content_type = r.headers.get('content-type', '')
            if 'image' not in content_type:
                print(f"❌ Contenido no es imagen ({content_type}): {img_url}")
                return None, None, None
                
            return r, True, None
            
        except requests.exceptions.Timeout:
            print(f"⏱️ Timeout al descargar: {img_url}")
            return None, None, "timeout"
        except requests.exceptions.ConnectionError:
            print(f"🌐 Error de conexión al descargar: {img_url}")
            return None, None, "connection"
        except Exception as e:
            print(f"🚨 Error inesperado al descargar {img_url}: {e}")
            return None, None, "error"

    while idx < len(paragraphs):
        par = paragraphs[idx]
        text = par.Range.Text
        m_pipe = pipe_pattern.search(text)
        if m_pipe:
            has_pipes = True
            match = m_pipe
        else:
            m_html = html_pattern.search(text)
            if m_html:
                has_pipes = False
                match = m_html
            else:
                idx += 1
                continue
                
        img_url = match.group(1)
        print(f"🔗 Procesando imagen: {img_url}")

        # Verificar conexión con la URL completa
        try:
            test_conn = requests.get(img_url, stream=True, timeout=5, headers=headers, allow_redirects=True)
            if test_conn.status_code != 200:
                print("Lo siento, la página está caída por lo cual no puedo descargar la imagen")
                idx += 1
                continue
            test_conn.close()
        except requests.RequestException:
            print("Lo siento, la página está caída por lo cual no puedo descargar la imagen")
            idx += 1
            continue

        # Descargar imagen si hay conexión
        print("Con respuesta, con gusto podré insertar la imagen en Microsoft Word y descargarla")
        response, success, error_type = descargar_imagen(img_url, session, headers)
        
        if not success:
            print(f"⏭️ Saltando imagen {img_url} debido a error de descarga")
            idx += 1
            continue
            
        try:
            # Generar nombre de archivo
            name = unquote(os.path.basename(img_url)).replace(' ','_')
            if not name.lower().endswith((".jpg",".jpeg",".png")): 
                name += ".jpg"
            path = os.path.join(attachments_dir, name)
            
            # Guardar imagen
            with open(path, 'wb') as f:
                for chunk in response.iter_content(1024): 
                    f.write(chunk)
            print(f"✅ Descargada: {name}")
            
            # Obtener dimensiones si PIL está disponible
            if pillow_installed:
                try:
                    with Image.open(path) as im: 
                        w, h = im.size
                except Exception as e:
                    print(f"⚠️ Error obteniendo dimensiones de {name}: {e}")
                    w, h = 1000, 800  # Valores por defecto
            else: 
                w, h = 1000, 800  # Valores por defecto sin PIL
            
            # Definir rango para inserción
            rng = par.Range.Duplicate
            rng.Start = par.Range.Start + match.start()
            rng.End = par.Range.Start + match.end()
            
            # Insertar imagen en el documento
            if has_pipes:
                if insert_with_pipes(rng, path, w, h):
                    print(f"🖼️ Insertada con pipes: {name}")
                else:
                    print(f"❌ Error insertando con pipes: {name}")
            else:
                if insert_without_pipes(rng, path):
                    print(f"🖼️ Insertada sin pipes: {name}")
                else:
                    print(f"❌ Error insertando sin pipes: {name}")
            
            # Refrescar lista de párrafos después de la inserción
            paragraphs = list(doc.Paragraphs)
            idx += 1
            
        except Exception as e:
            print(f"🚨 Error procesando imagen {img_url}: {e}")
            idx += 1
            
    session.close()
    print("🏁 Procesamiento de imágenes completado")

# Ejecutar
procesar_imagenes_html(doc)





print("Imágenes con enlaces html página postimages.org postimg.cc")

def procesar_imagenes_html(doc):
    # PIL para medir orientación
    try:
        from PIL import Image
        pillow = True
    except ImportError:
        pillow = False

    # Carpeta de attachments
    script_dir = os.path.dirname(os.path.abspath(__file__))
    attach_dir = os.path.join(script_dir, ".attachments")
    os.makedirs(attach_dir, exist_ok=True)

    # Patrones para HTML postimg.cc con y sin pipes
    pipe_pat = re.compile(
        r"\|\(a\s+[^)]+\)\(img\s+[^)]*src=['\\"]([^'\\"]+)['\\"][^)]*\)\(/a\)\|",
        re.IGNORECASE
    )
    html_pat = re.compile(
        r"\(a\s+[^)]+\)\(img\s+[^)]*src=['\\"]([^'\\"]+)['\\"][^)]*\)\(/a\)",
        re.IGNORECASE
    )

    # Constantes de tamaño
    MAX_H    = 6 * 28.3465   # 6 cm en puntos
    PIPE_SZ  = 12            # tamaño de "|" en pts
    PIPE_H   = PIPE_SZ + 2   # altura de fila de pipes

    def insert_with_pipes(rng, img_path, w, h):
        try:
            rng.Delete()
            rng.Collapse(win32.constants.wdCollapseStart)

            MAX_H    = 6 * 28.3465
            MAX_W_PT = 15 * 28.3465
            PIPE_SZ  = 12
            PIPE_H   = PIPE_SZ + 2

            # Tabla para medir imagen
            temp = doc.Tables.Add(rng, 1, 1)
            temp.Borders.Enable = False
            cell = temp.Cell(1, 1)
            pic  = cell.Range.InlineShapes.AddPicture(img_path, False, True)
            if pic.Height > MAX_H:
                pic.Height = MAX_H
            img_width = pic.Width
            temp.Delete()

            if img_width < MAX_W_PT:
                # Tabla 3x2: (fila superior vacía, fila imagen combinada, fila inferior con dos pipes)
                tbl = doc.Tables.Add(rng, 3, 2)
                tbl.Borders.Enable = False
                tbl.Rows.AllowBreakAcrossPages = False

                # Fila 1: vacía
                tbl.Cell(1, 1).Range.Text = ""
                tbl.Cell(1, 2).Range.Text = ""

                # Fila 2: imagen centrada en celda combinada
                tbl.Cell(2, 1).Merge(tbl.Cell(2, 2))
                img_cell = tbl.Cell(2, 1)
                pic = img_cell.Range.InlineShapes.AddPicture(img_path, False, True)
                if pic.Height > MAX_H:
                    pic.Height = MAX_H
                img_cell.Range.ParagraphFormat.Alignment = win32.constants.wdAlignParagraphCenter

                # Fila 3: pipes en celdas izquierda y derecha
                left_pipe  = tbl.Cell(3, 1)
                right_pipe = tbl.Cell(3, 2)

                left_pipe.Range.Text = "|"
                left_pipe.Range.Font.Size = PIPE_SZ
                left_pipe.Range.ParagraphFormat.Alignment = win32.constants.wdAlignParagraphLeft

                right_pipe.Range.Text = "|"
                right_pipe.Range.Font.Size = PIPE_SZ
                right_pipe.Range.ParagraphFormat.Alignment = win32.constants.wdAlignParagraphRight

                # Ajustar altura fila inferior
                tbl.Rows(3).HeightRule = win32.constants.wdRowHeightExactly
                tbl.Rows(3).Height = PIPE_H

            else:
                # Imagen ancha: tabla 3x1 con pipe arriba y abajo
                tbl = doc.Tables.Add(rng, 3, 1)
                tbl.Borders.Enable = False

                tbl.Cell(1, 1).Range.Text = "|"
                tbl.Cell(1, 1).Range.Font.Size = PIPE_SZ
                tbl.Cell(1, 1).Range.ParagraphFormat.Alignment = win32.constants.wdAlignParagraphLeft

                mid_cell = tbl.Cell(2, 1)
                pic = mid_cell.Range.InlineShapes.AddPicture(img_path, False, True)
                if pic.Height > MAX_H:
                    pic.Height = MAX_H
                mid_cell.Range.ParagraphFormat.Alignment = win32.constants.wdAlignParagraphCenter

                tbl.Cell(3, 1).Range.Text = "|"
                tbl.Cell(3, 1).Range.Font.Size = PIPE_SZ
                tbl.Cell(3, 1).Range.ParagraphFormat.Alignment = win32.constants.wdAlignParagraphLeft

                tbl.Rows(1).HeightRule = win32.constants.wdRowHeightExactly
                tbl.Rows(1).Height = PIPE_H
                tbl.Rows(3).HeightRule = win32.constants.wdRowHeightExactly
                tbl.Rows(3).Height = PIPE_H

            tbl.Columns.AutoFit()
            return True
        except Exception as e:
            print(f"❌ Error insert_with_pipes: {e}")
            return False


    def insert_without_pipes(rng, img_path):
        try:
            rng.Delete(); rng.Collapse(win32.constants.wdCollapseStart)
            pic = rng.InlineShapes.AddPicture(img_path, False, True)
            if pic.Height > MAX_H: pic.Height = MAX_H
            rng.ParagraphFormat.Alignment = win32.constants.wdAlignParagraphCenter
            return True
        except Exception as e:
            print(f"Error insert_without_pipes: {e}")
            return False

    # Recorremos párrafos
    for para in list(doc.Paragraphs):
        text = para.Range.Text
        m = pipe_pat.search(text)
        has_pipes = bool(m)
        if not m:
            m = html_pat.search(text)
            if not m:
                continue

        url = m.group(1)
        s, e = m.span()
        print(f"Procesando {url} (pipes={'sí' if has_pipes else 'no'})")

        # Verificar conexión con la URL completa
        headers = {'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'}
        try:
            # Usamos GET con stream para verificar, pero sin descargar todo
            test_conn = requests.get(url, stream=True, timeout=5, headers=headers, allow_redirects=True)
            if test_conn.status_code != 200:
                print("Lo siento, la página está caída por lo cual no puedo descargar la imagen")
                continue
            # Cerramos la conexión de prueba para no consumir recursos
            test_conn.close()
        except requests.RequestException:
            print("Lo siento, la página está caída por lo cual no puedo descargar la imagen")
            continue

        # Descargar imagen si hay conexión
        print("Con respuesta, con gusto podré insertar la imagen en Microsoft Word y descargarla")
        resp = requests.get(url, stream=True, headers=headers)
        if resp.status_code != 200 or 'image' not in resp.headers.get('content-type',''):
            print(f"Error descarga {resp.status_code}")
            continue

        # Guardar
        fn = unquote(os.path.basename(url)).replace(' ','_')
        if not fn.lower().endswith((".jpg",".jpeg",".png")): fn += ".jpg"
        path = os.path.join(attach_dir, fn)
        with open(path, 'wb') as f:
            for chunk in resp.iter_content(1024):
                f.write(chunk)

        # Medir
        if pillow:
            from PIL import Image
            w,h = Image.open(path).size
        else:
            w,h = 1000,0

        # Rango a eliminar
        rng = para.Range.Duplicate
        rng.Start = para.Range.Start + s
        rng.End   = para.Range.Start + e

        # Insertar
        if has_pipes:
            insert_with_pipes(rng, path, w, h)
        else:
            insert_without_pipes(rng, path)

        print(f"Insertada: {fn} (pipes={'sí' if has_pipes else 'no'})")

# Al final
procesar_imagenes_html(doc)





print("Descargar imágenes HTML formato específico de la página www.photo-pick.com (con detección de pipes en enlace)")

# Patrones: con y sin pipes
pipe_pattern = re.compile(r"\|https://www\.photo-pick\.com/online/([A-Z0-9]+)\.link\|", re.IGNORECASE)
link_pattern = re.compile(r"https://www\.photo-pick\.com/online/([A-Z0-9]+)\.link", re.IGNORECASE)

# Funciones auxiliares

def insert_image_with_pipes(doc, range_obj, image_path):
    try:
        range_obj.Delete()
        range_obj.Collapse(win32.constants.wdCollapseStart)
        max_h = 6 * 28.3465
        max_w_cm = 15
        max_w = max_w_cm * 28.3465
        # medir temporal
        temp = range_obj.Duplicate
        tbl = doc.Tables.Add(temp, 1, 1)
        tbl.Borders.Enable = False
        pic = tbl.Cell(1,1).Range.InlineShapes.AddPicture(image_path, False, True)
        if pic.Height > max_h: pic.Height = max_h
        img_w = pic.Width
        tbl.Delete()
        # crear tabla final
        if img_w < max_w:
            table = doc.Tables.Add(range_obj, 3, 1)
            table.Borders.Enable = False
            table.Rows(1).Height = 1
            table.Cell(1,1).Range.Text = ""
            cell_img = table.Cell(2,1)
            pic = cell_img.Range.InlineShapes.AddPicture(image_path, False, True)
            if pic.Height > max_h: pic.Height = max_h
            cell_img.Range.ParagraphFormat.Alignment = win32.constants.wdAlignParagraphCenter
            table.Cell(3,1).Split(1,2)
            table.PreferredWidth = img_w
            for c in (table.Cell(3,1), table.Cell(3,2)):
                c.Width = img_w/2 - 1
                c.Range.ParagraphFormat.LeftIndent = 0
                c.Range.ParagraphFormat.RightIndent = 0
                c.Range.Text = "|"
                c.Range.ParagraphFormat.Alignment = win32.constants.wdAlignParagraphLeft
        else:
            table = doc.Tables.Add(range_obj, 3, 1)
            table.Borders.Enable = False
            top = table.Cell(1,1)
            top.Range.Text = "|"
            top.Range.ParagraphFormat.Alignment = win32.constants.wdAlignParagraphLeft
            cell_img = table.Cell(2,1)
            pic = cell_img.Range.InlineShapes.AddPicture(image_path, False, True)
            if pic.Height > max_h: pic.Height = max_h
            cell_img.Range.ParagraphFormat.Alignment = win32.constants.wdAlignParagraphCenter
            bot = table.Cell(3,1)
            bot.Range.Text = "|"
            bot.Range.ParagraphFormat.Alignment = win32.constants.wdAlignParagraphLeft
        table.Columns.AutoFit()
        return True
    except Exception as e:
        print(f"Error insert_with_pipes: {e}")
        return False


def insert_image_without_pipes(doc, range_obj, image_path):
    try:
        range_obj.Delete()
        range_obj.Collapse(win32.constants.wdCollapseStart)
        pic = range_obj.InlineShapes.AddPicture(image_path, False, True)
        max_h = 6 * 28.3465
        if pic.Height > max_h: pic.Height = max_h
        range_obj.ParagraphFormat.Alignment = win32.constants.wdAlignParagraphCenter
        return True
    except Exception as e:
        print(f"Error insert_without_pipes: {e}")
        return False


def procesar_imagenes_photopick(doc):
    script_dir = os.path.dirname(os.path.abspath(__file__))
    attach_dir = os.path.join(script_dir, ".attachments")
    os.makedirs(attach_dir, exist_ok=True)
    print(f"📂 Directorio .attachments: {attach_dir}")

    paragraphs = list(doc.Paragraphs)
    session = requests.Session()
    headers = {'User-Agent': 'Mozilla/5.0'}
    idx = 0
    while idx < len(paragraphs):
        par = paragraphs[idx]
        txt = par.Range.Text
        m_pipe = pipe_pattern.search(txt)
        m_link = link_pattern.search(txt)
        if m_pipe:
            has_pipes = True
            photo_id = m_pipe.group(1)
            span = m_pipe.span()
        elif m_link:
            has_pipes = False
            photo_id = m_link.group(1)
            span = m_link.span()
        else:
            idx += 1
            continue
        url_page = f"https://www.photo-pick.com/online/{photo_id}.link"
        try:
            r = session.get(url_page, headers=headers, timeout=15)
            if r.status_code != 200:
                raise ValueError(f"HTTP {r.status_code}")
            soup = BeautifulSoup(r.text, 'html.parser')
            meta = soup.find('meta', attrs={'property':'og:image'})
            if not meta or not meta.get('content'):
                raise ValueError("og:image no encontrado")
            url_img = meta['content']
            resp = session.get(url_img, headers=headers, stream=True, timeout=15)
            if resp.status_code != 200 or 'image' not in resp.headers.get('content-type',''):
                raise ValueError("Descarga imagen fallida")
            fname = unquote(os.path.basename(url_img)).replace(' ','_')
            if not fname.lower().endswith((".jpg",".jpeg",".png")):
                fname += ".jpg"
            path = os.path.join(attach_dir, fname)
            with open(path, 'wb') as f:
                for ch in resp.iter_content(1024): f.write(ch)
            print(f"✅ Descargada: {fname}")
            # definir rango de match
            full = par.Range.Duplicate
            full.Start = par.Range.Start + span[0]
            full.End = par.Range.Start + span[1]
            # insertar según pipes
            if has_pipes:
                ok = insert_image_with_pipes(doc, full, path)
            else:
                ok = insert_image_without_pipes(doc, full, path)
            print(f"🖼️ Insertada {'con pipes' if has_pipes else 'sin pipes'}: {fname}")
            # refrescar y avanzar
            paragraphs = list(doc.Paragraphs)
            idx += 1
        except Exception as e:
            print(f"🚨 Error procesar {url_page}: {e}")
            idx += 1
    session.close()
    print("✅ Proceso terminado.")

# Ejecutar
procesar_imagenes_photopick(doc)






                

print("Convirtiendo enlaces Markdown a hipervínculos de Word...")

# Expresiones regulares para distintos tipos de enlaces
markdown_pattern = r"(?<!\!)\[(?!https?://)(.*?)\]\((?!url:)(.*?)\)"
url_pattern = r"(https?://[^\s\)]+)"
markdown_http_pattern = r"\[(https?://[^\]]+?)\]\(\)"  # Nueva expresión regular

# Variables para detección de bloques de código
in_codeblock = False

# Recorrer todos los párrafos
for i in range(1, doc.Paragraphs.Count + 1):
    paragraph = doc.Paragraphs.Item(i)
    paragraph_text = paragraph.Range.Text

    # Verificar si hemos encontrado el inicio o fin de un bloque de código con "```"
    if '```' in paragraph_text:
        in_codeblock = not in_codeblock
        continue  # Si es un delimitador de bloque de código, lo ignoramos y pasamos al siguiente párrafo

    # Si estamos dentro de un bloque de código, no procesamos el texto
    if in_codeblock:
        continue

    # Procesar el nuevo tipo de enlace `[http://loquesea]()`
    http_matches = list(re.finditer(markdown_http_pattern, paragraph_text))
    if http_matches:
        for match in reversed(http_matches):
            url = match.group(1).strip()

            # Posición en el texto
            start = paragraph.Range.Start + match.start()
            end = paragraph.Range.Start + match.end()

            # Validar el rango
            if start < paragraph.Range.End:
                match_range = doc.Range(Start=start, End=end)
                try:
                    # Reemplazar el enlace Markdown con la URL como texto visible
                    match_range.Text = url
                    # Actualizar el rango para abarcar solo el nuevo texto
                    match_range = doc.Range(Start=start, End=start + len(url))
                    # Añadir el hipervínculo
                    hyperlink = doc.Hyperlinks.Add(Anchor=match_range, Address=url, TextToDisplay=url)
                    # Aplicar formato azul y subrayado al rango del hipervínculo
                    hyperlink.Range.Font.ColorIndex = constants.wdBlue
                    hyperlink.Range.Font.Underline = constants.wdUnderlineSingle
                except Exception as e:
                    print(f"Error añadiendo hipervínculo para '{url}': {e}")

    # Procesar enlaces Markdown regulares
    matches = list(re.finditer(markdown_pattern, paragraph_text))
    if matches:
        for match in reversed(matches):
            display_text = match.group(1).strip()
            url = match.group(2).strip()
            if not url:
                url = display_text
            start = paragraph.Range.Start + match.start()
            end = paragraph.Range.Start + match.end()
            if start < paragraph.Range.End:
                match_range = doc.Range(Start=start, End=end)
                try:
                    match_range.Text = display_text
                    match_range = doc.Range(Start=start, End=start + len(display_text))
                    hyperlink = doc.Hyperlinks.Add(Anchor=match_range, Address=url, TextToDisplay=display_text)
                    hyperlink.Range.Font.ColorIndex = constants.wdBlue
                    hyperlink.Range.Font.Underline = constants.wdUnderlineSingle
                except Exception as e:
                    print(f"Error añadiendo hipervínculo para '{display_text}': {e}")

    # Procesar URLs planas
    url_matches = list(re.finditer(url_pattern, paragraph_text))
    if url_matches:
        for match in reversed(url_matches):
            url = match.group(0)

            # Limpieza de paréntesis o corchetes al final del enlace
            url = url.rstrip("])(")

            display_text = url
            start = paragraph.Range.Start + match.start()
            end = paragraph.Range.Start + match.end()
            if start < paragraph.Range.End:
                match_range = doc.Range(Start=start, End=end)
                try:
                    hyperlink = doc.Hyperlinks.Add(Anchor=match_range, Address=url, TextToDisplay=display_text)
                    hyperlink.Range.Font.ColorIndex = constants.wdBlue
                    hyperlink.Range.Font.Underline = constants.wdUnderlineSingle
                except Exception as e:
                    print(f"Error añadiendo hipervínculo para '{url}': {e}")




print("aplicar código markdown normal")
# Función para crear colores RGB
def RGB(r, g, b):
    return r + (g << 8) + (b << 16)

# Variables para detección de bloques de código
found_codeblock = False
in_codeblock = False

# Expresión regular para detectar "negrit```" o "```negrit"
pattern = re.compile(r'(negrit```|```negrit)')

# Iteración sobre párrafos
for paragraph in doc.Paragraphs:
    text = paragraph.Range.Text

    # Asegurarse de que el texto es una cadena
    if not isinstance(text, str):
        continue  # Si no es una cadena, pasa al siguiente párrafo

    # Verificar si hemos encontrado el inicio de un bloque de código con exactamente tres acentos invertidos
    if text.count('```') == 1 and text.count('`') == 3:
        # Verificar si el texto contiene "negrit```" o "```negrit"
        if re.search(pattern, text):
            # Reemplazar ``` con ````
            updated_text = text.replace('```', '````')
            paragraph.Range.Text = updated_text
        else:
            in_codeblock = not in_codeblock
            found_codeblock = True
            
            # Eliminar el párrafo que contiene "```"
            paragraph.Range.Delete()
            continue  # Saltar el delimitador

    # Si estamos dentro de un bloque de código, cambiar la fuente y el color del texto y del fondo
    if in_codeblock:
        # Cambiar la fuente, tamaño y formato
        paragraph.Range.Font.Name = "Consolas"
        paragraph.Range.Font.Size = 11
        paragraph.Range.Font.Bold = False  # Sin negrita
        paragraph.Range.Font.Color = RGB(255, 255, 255)  # Blanco
        paragraph.Range.Shading.BackgroundPatternColor = RGB(0, 0, 0)  # Negro


    # Si estamos dentro o fuera del bloque, eliminar los caracteres ``` antes y después, ignorando si hay "negrit"
    if text.count('```') == 1 and not in_codeblock:
        # Si se detecta la palabra "negrit", no cambiar el texto
        if re.search(pattern, text):
            continue
        updated_text = text.replace('```', '')
        paragraph.Range.Text = updated_text

# Imprimir un mensaje si no se encontró ningún bloque de código
if not found_codeblock:
    print("No se encontró ningún bloque de código en el documento.")


print("bloque código especial negrit y etiquetas (code)")

# Variables para detección de bloques de código
found_codeblock = False
in_codeblock = False
in_code_tag = False

# Expresión regular para detectar "negrit````" o "````negrit"
pattern_codeblock2 = re.compile(r'(negrit````|````negrit)')
# Expresión regular para detectar las etiquetas (code) y (/code)
pattern_code_tag_start = re.compile(r'\(code\)')
pattern_code_tag_end = re.compile(r'\(/code\)')

# Iteración sobre párrafos
for paragraph in doc.Paragraphs:
    text = paragraph.Range.Text.strip()  # Eliminar espacios y saltos de línea del principio y final

    # Verificar si hemos encontrado el inicio de un bloque de código con "````" y "negrit"
    if text.count('````') == 1 and text.count('`') == 4:
        if re.search(pattern_codeblock2, text):
            in_codeblock = not in_codeblock
            found_codeblock = True

            # Aplicar formato de bloque de código especial
            paragraph.Range.Font.Name = "Consolas"
            paragraph.Range.Font.Bold = True
            paragraph.Range.Font.Color = win32api.RGB(0, 0, 0)  # Texto en color negro
            paragraph.Range.Shading.BackgroundPatternColor = win32api.RGB(192, 192, 192)  # Fondo gris

            # **Ajustar espaciado superior**
            paragraph.SpaceBefore = 0  # Espaciado antes del párrafo
            paragraph.SpaceAfter = 0   # Espaciado después del párrafo
            paragraph.LineSpacingRule = 1  # Espaciado de líneas simple

            # Eliminar las etiquetas "negrit" y "````"
            updated_text = re.sub(r'negrit````|````negrit', '', text)
            paragraph.Range.Text = updated_text

    # Detectar el inicio de un bloque con etiquetas (code)
    if re.search(pattern_code_tag_start, text):
        in_code_tag = True
        found_codeblock = True

        # Aplicar formato de bloque de código
        paragraph.Range.Font.Name = "Consolas"
        paragraph.Range.Font.Bold = True
        paragraph.Range.Font.Color = win32api.RGB(0, 0, 0)  # Texto en color negro
        paragraph.Range.Shading.BackgroundPatternColor = win32api.RGB(192, 192, 192)  # Fondo gris

        # **Ajustar espaciado superior**
        paragraph.SpaceBefore = 0  # Espaciado antes del párrafo
        paragraph.SpaceAfter = 0   # Espaciado después del párrafo
        paragraph.LineSpacingRule = 1  # Espaciado de líneas simple

        # Eliminar la etiqueta (code)
        updated_text = re.sub(r'\(code\)', '', text)
        paragraph.Range.Text = updated_text

    # Detectar el final de un bloque con etiquetas (/code)
    if re.search(pattern_code_tag_end, text):
        in_code_tag = False

        # Eliminar la etiqueta (/code)
        updated_text = re.sub(r'\(/code\)', '', text)
        paragraph.Range.Text = updated_text

    # Si estamos dentro de un bloque (negrit o (code)), aplicar el formato
    if in_codeblock or in_code_tag:
        paragraph.Range.Font.Name = "Consolas"
        paragraph.Range.Font.Bold = True
        paragraph.Range.Font.Color = win32api.RGB(0, 0, 0)  # Texto en color negro
        paragraph.Range.Shading.BackgroundPatternColor = win32api.RGB(192, 192, 192)  # Fondo gris

        # **Ajustar espaciado superior**
        paragraph.SpaceBefore = 0  # Espaciado antes del párrafo
        paragraph.SpaceAfter = 0   # Espaciado después del párrafo
        paragraph.LineSpacingRule = 1  # Espaciado de líneas simple

    # Si estamos fuera de los bloques, eliminar los "````" en el texto
    if text.count('````') == 1 and not in_codeblock:
        if not re.search(pattern_codeblock2, text):
            updated_text = text.replace('````', '')
            paragraph.Range.Text = updated_text

# Imprimir un mensaje si no se encontró ningún bloque de código
if not found_codeblock:
    print("No se encontró ningún bloque de código especial con negrit o etiquetas (code).")






print("hypervinculo this link after x in red")

# Definir el patrón de búsqueda para enlaces en formato Markdown y formato especial con 'url:'
markdown_pattern2 = re.compile(r'\[([^\]]+)\]\(url:\s*[\\'"]([^\\'"]+)[\\'"]\)')

# Iterar sobre todos los párrafos en el documento
for paragraph in doc.Paragraphs:
    match = markdown_pattern2.search(paragraph.Range.Text)
    if match:
        # Extraemos el texto del enlace y la URL
        link_text = match.group(1)  # El texto que se muestra en el enlace
        url = match.group(2)        # La URL del enlace

        # Creamos un nuevo objeto de rango que contiene el texto del enlace
        link_range = doc.Range(paragraph.Range.Start, paragraph.Range.End)
        
        # Reemplazar el texto en el rango con solo el texto del enlace (sin la URL ni la palabra 'url:')
        link_range.Text = re.sub(markdown_pattern2, link_text, link_range.Text)

        # Agregamos el hipervínculo de Word
        hyperlink = doc.Hyperlinks.Add(link_range, url, TextToDisplay=link_text)

        # Cambiar el color del texto del hipervínculo a rojo (en formato BGR)
        red_color_bgr = 255 | (0 << 8) | (0 << 16)
        hyperlink.Range.Font.Color = red_color_bgr
        hyperlink.Range.Font.Underline = True

        # Agregar una "x" al final del enlace
        hyperlink.Range.InsertAfter(' x')
      
   

   

# Iterar sobre cada párrafo en el documento
print("Applying correct good formatting style for exact '---' lines, outside code blocks")

in_codeblock = False  # Rastrea si estamos dentro de ```bloques de código```

for paragraph in doc.Paragraphs:
    text = paragraph.Range.Text

    # 1) Detectar toggles de bloque de código (líneas con exactamente 3 backticks)
    if text.count('```') == 1 and text.count('`') == 3:
        in_codeblock = not in_codeblock
        # No procesamos más esta línea
        continue

    # 2) Si estamos dentro de un bloque de código, no hacemos nada
    if in_codeblock:
        continue

    # 3) Fuera de código: solo actuar si el párrafo es EXACTAMENTE '---'
    if text.strip() == '---':
        # 3.1) Dejar el párrafo “vacío” (o eliminarlo si lo prefieres)
        paragraph.Range.Text = ''

        # 3.2) Aplicar borde superior
        border = paragraph.Range.Borders(win32.constants.wdBorderTop)
        border.LineStyle  = win32.constants.wdLineStyleSingle
        border.LineWidth  = win32.constants.wdLineWidth050pt
        border.Color      = win32api.RGB(234, 234, 234)


        
        

print("Aplicar color amarillo al texto rodeado por dos acentos graves y eliminar los acentos graves")
try:
    find_object = doc.Content.Find
    find_object.ClearFormatting()
    print("Depuración: Formato limpiado.")

    # Buscar texto rodeado por dos acentos graves en el contenido principal
    find_object.Text = "``*``"  # Usa comodín para encontrar texto entre dos acentos graves
    find_object.MatchWildcards = True

    # Aplicar color amarillo al texto encontrado
    while find_object.Execute():
        found_range = find_object.Parent
        found_range.Font.Shading.BackgroundPatternColor = 65535  # Color amarillo
        found_range.Text = found_range.Text[2:-2]  # Eliminar los acentos graves

    # Buscar y procesar texto en cada celda de cada tabla
    for table in doc.Tables:
        for row in table.Rows:
            for cell in row.Cells:
                find_object = cell.Range.Find
                find_object.ClearFormatting()
                find_object.Text = "``*``"  # Usa comodín para encontrar texto entre dos acentos graves
                find_object.MatchWildcards = True
                while find_object.Execute():
                    found_range = find_object.Parent
                    found_range.Font.Shading.BackgroundPatternColor = 65535  # Color amarillo
                    found_range.Text = found_range.Text[2:-2]  # Eliminar los acentos graves

    print("Depuración: Texto resaltado en color amarillo y acentos graves eliminados.")

except Exception as e:
    print(f"Error al aplicar el color amarillo o eliminar los acentos graves: {e}")



print("Aplicar color sombreado gris al texto rodeado por un acento grave y eliminar el acento grave")
# Aplicar color sombreado gris al texto rodeado por un acento grave y eliminar el acento grave
try:
    find_object = doc.Content.Find
    find_object.ClearFormatting()
    print("Depuración: Formato limpiado.")

    # Buscar texto rodeado por un acento grave en el contenido principal
    find_object.Text = "`*`"  # Usa comodín para encontrar texto entre un acento grave
    find_object.MatchWildcards = True

    # Aplicar color sombreado gris al texto encontrado
    while find_object.Execute():
        found_range = find_object.Parent
        found_range.Font.Shading.BackgroundPatternColor = 12632256  # Color sombreado gris
        found_range.Text = found_range.Text[1:-1]  # Eliminar el acento grave

    # Buscar y procesar texto en cada celda de cada tabla
    for table in doc.Tables:
        for row in table.Rows:
            for cell in row.Cells:
                find_object = cell.Range.Find
                find_object.ClearFormatting()
                find_object.Text = "`*`"  # Usa comodín para encontrar texto entre un acento grave
                find_object.MatchWildcards = True
                while find_object.Execute():
                    found_range = find_object.Parent
                    found_range.Font.Shading.BackgroundPatternColor = 12632256  # Color sombreado gris
                    found_range.Text = found_range.Text[1:-1]  # Eliminar el acento grave

    print("Depuración: Texto resaltado en color gris y acentos graves eliminados.")

except Exception as e:
    print(f"Error al aplicar el color gris o eliminar los acentos graves: {e}")




print("Aplicar Fit Text y Wrap Text a celdas con texto mayormente sombreado (gris o amarillo) y con más de un espacio")
try:
    # Procesar cada tabla en el documento
    for table in doc.Tables:
        for row in table.Rows:
            for cell in row.Cells:
                # Obtener el rango de texto en la celda
                range_in_cell = cell.Range
                # Filtrar caracteres relevantes (excluir fin de celda y espacios en blanco al final)
                relevant_characters = [
                    character for character in range_in_cell.Characters
                    if character.Text.strip()  # Ignorar caracteres vacíos o de fin de celda
                ]

                # Verificar si la mayoría de los caracteres relevantes están sombreados en gris o amarillo
                shaded_characters = sum(
                    1 for character in relevant_characters
                    if character.Font.Shading.BackgroundPatternColor in [12632256, 65535]  # Gris o amarillo
                )
                total_relevant_characters = len(relevant_characters)
                shading_ratio = shaded_characters / total_relevant_characters if total_relevant_characters > 0 else 0

                # Considerar 'mayormente sombreado' si más del 90% de los caracteres están sombreados
                is_majority_shaded = shading_ratio > 0.9

                # Contar la cantidad de espacios en el texto
                space_count = range_in_cell.Text.count(' ')

                # Verificar si el texto cumple las condiciones
                if is_majority_shaded and space_count > 1:
                    cell.FitText = True
                    cell.WordWrap = True
    print("Proceso completado: Propiedades aplicadas a celdas tanto fit como wrap.")
except Exception as e:
    print(f"Error al aplicar propiedades a las celdas: {e}")



print("aplicar sombreado gris dentro de listas y viñetas")
def apply_gray_shading_to_lists(doc):
    try:
        find_object = doc.Content.Find
        find_object.ClearFormatting()
        print("Depuración: Formato limpiado.")

        # Buscar texto rodeado por un acento grave en el contenido principal
        find_object.Text = "`*`"  # Usa comodín para encontrar texto entre un acento grave
        find_object.MatchWildcards = True

        # Aplicar color sombreado gris al texto encontrado
        while find_object.Execute():
            found_range = find_object.Parent
            found_range.Font.Shading.BackgroundPatternColor = 12632256  # Color sombreado gris
            found_range.Text = found_range.Text[1:-1]  # Eliminar el acento grave

        # Buscar y procesar texto en cada párrafo de las listas
        for paragraph in doc.Paragraphs:
            find_object = paragraph.Range.Find
            find_object.ClearFormatting()
            find_object.Text = "`*`"
            find_object.MatchWildcards = True
            while find_object.Execute():
                found_range = find_object.Parent
                found_range.Font.Shading.BackgroundPatternColor = 12632256
                found_range.Text = found_range.Text[1:-1]

        print("Depuración: Texto resaltado en color gris y acentos graves eliminados.")

    except Exception as e:
        print(f"Error al aplicar el color gris o eliminar los acentos graves: {e}")
        
apply_gray_shading_to_lists(doc)
        
        

print("html tag automático y rápido")

# Diccionario de colores con valores RGB
colores = {
    'azul': win32api.RGB(0, 0, 255),
    'blue': win32api.RGB(0, 0, 255),
    'yellow': win32api.RGB(255, 255, 0),
    'amarillo': win32api.RGB(255, 255, 0),
    'verde': win32api.RGB(0, 128, 0),
    'green': win32api.RGB(0, 128, 0),
    'marrón': win32api.RGB(165, 42, 42),
    'brown': win32api.RGB(165, 42, 42),
    'rosa': win32api.RGB(255, 105, 180),
    'pink': win32api.RGB(255, 105, 180),
    'rojo': win32api.RGB(255, 0, 0),
    'red': win32api.RGB(255, 0, 0),
    'crimson': win32api.RGB(220, 20, 60),
    'teal': win32api.RGB(0, 128, 128),
    'purple': win32api.RGB(128, 0, 128),
    'violeta': win32api.RGB(128, 0, 128),
    'colour': win32api.RGB(0, 0, 0),
    'black': win32api.RGB(0, 0, 0)
}

def aplicar_formatos(paragraph):
    # Expresión regular para etiquetas (span style=color:color) con (b)...(/b)
    pattern_color_bold = re.compile(r'\(b\)\(span\s+style\s*=\s*["\\']?color\s*:\s*([a-z]+)\s*["\\']?\)([^()]*?)\(/b\)\(/span\)', re.IGNORECASE | re.DOTALL)
    # Expresión regular para etiquetas (span style=color:color) sin (b)...(/b)
    pattern_color = re.compile(r'\(span\s+style\s*=\s*["\\']?color\s*:\s*([a-z]+)\s*["\\']?\)([^()]*?)\(/span\)', re.IGNORECASE | re.DOTALL)

    paragraph_text = paragraph.Range.Text

    # Procesar coincidencias de etiquetas con (b)...(/b)
    for match in pattern_color_bold.finditer(paragraph_text):
        color = match.group(1).strip().lower() if match.group(1) else 'black'
        content = match.group(2) if match.group(2) else ''

        if color in colores:
            try:
                # Calcular posiciones en el documento para la parte a colorear y poner en negrita
                start = paragraph.Range.Start + match.start(2)
                end = start + len(content)
                content_range = doc.Range(start, end)
                content_range.Font.Color = colores[color]  # Aplicar el color
                content_range.Font.Bold = True  # Aplicar negrita

            except Exception as e:
                print(f"No se pudo cambiar el color y negrita: {e}")
        else:
            print(f"Color desconocido: {color}")

    # Procesar coincidencias de etiquetas sin (b)...(/b)
    for match in pattern_color.finditer(paragraph_text):
        color = match.group(1).strip().lower() if match.group(1) else 'black'
        content = match.group(2) if match.group(2) else ''

        if color in colores:
            try:
                # Calcular posiciones en el documento para la parte a colorear
                start = paragraph.Range.Start + match.start(2)
                end = start + len(content)
                content_range = doc.Range(start, end)
                content_range.Font.Color = colores[color]  # Aplicar el color

            except Exception as e:
                print(f"No se pudo cambiar el color: {e}")
        else:
            print(f"Color desconocido: {color}")

    # Eliminar todas las etiquetas de apertura y cierre al final
    eliminar_etiquetas(paragraph.Range)

def eliminar_etiquetas(rango):
    # Eliminar etiquetas de cierre
    rango.Find.ClearFormatting()
    rango.Find.Replacement.ClearFormatting()
    rango.Find.Execute(FindText=r"\(/span\)", ReplaceWith="", Replace=win32com.client.constants.wdReplaceAll, Forward=True, MatchWildcards=True)
    rango.Find.Execute(FindText=r"\(/b\)", ReplaceWith="", Replace=win32com.client.constants.wdReplaceAll, Forward=True, MatchWildcards=True)

    # Eliminar etiquetas de apertura
    rango.Find.Execute(FindText=r"\(b\)", ReplaceWith="", Replace=win32com.client.constants.wdReplaceAll, Forward=True, MatchWildcards=True)
    rango.Find.Execute(FindText=r"\(span[!)]@\)", ReplaceWith="", Replace=win32com.client.constants.wdReplaceAll, Forward=True, MatchWildcards=True)

def process_document():
    for paragraph in doc.Paragraphs:
        # Limitar a párrafos que contienen etiquetas (span style="color")
        if re.search(r'\(span\s+style\s*=\s*["\\']?color\s*:\s*[a-z]+\s*["\\']?\)', paragraph.Range.Text, re.IGNORECASE):
            aplicar_formatos(paragraph)

# Ejecuta el procesamiento del documento
process_document()



# Patrón regex para identificar URLs
url_pattern = re.compile(r'\b((http|https):\/\/)?[^\s()<>]+(?:\.[a-z]{2,})')

# Patrón regex para identificar imágenes en Markdown
image_pattern = re.compile(r'!\[([^\]]*)\]\(([^)]+)\)')

# Iterar sobre todos los párrafos en el documento
for paragraph in doc.Paragraphs:
    # Buscar todas las URLs en el párrafo
    for match in re.finditer(url_pattern, paragraph.Range.Text):
        # Obtener la URL y su posición en el texto del párrafo
        url = match.group()
        start_pos = match.start()
        end_pos = match.end()

        # Comprobar si la URL es parte de una imagen en Markdown
        if image_pattern.search(paragraph.Range.Text[start_pos:end_pos]):
            print(f"Ignoring the URL '{url}' because it is part of an image description in Markdown.")
            continue  # Si es así, ignora la URL y pasa a la siguiente

        # Crear un rango que solo incluye la URL y verificar si es válido
        url_range = doc.Range(paragraph.Range.Start + start_pos, paragraph.Range.Start + end_pos)
        
        if url_range.Start < url_range.End:  # Verificar si el rango es válido antes de añadir el hipervínculo
            try:
                doc.Hyperlinks.Add(Anchor=url_range, Address=url)
            except Exception as e:
                print(f"Could not add hyperlink for URL '{url}': {e}")
        else:
            print(f"Invalid range for URL '{url}'. Skipping hyperlink addition.")

# Itera sobre todos los campos en el documento
for field in doc.Fields:
    try:
        # Comprueba si el campo es un hipervínculo
        if field.Type == win32com.client.constants.wdFieldHyperlink:
            print("Found a hyperlink.")
            # Comprueba si el campo es parte de la tabla de contenido
            if field.Code.Text.startswith(" TOC "):
                print("The hyperlink is part of the table of contents, the formatting is not changed.")
                continue  # Si es parte de la tabla de contenido, no cambies el formato
            # Cambia el color del texto a azul
            field.Result.Font.Color = win32com.client.constants.wdColorBlue
            # Cambia el color del subrayado a azul
            field.Result.Font.UnderlineColor = win32com.client.constants.wdColorBlue
            # Aplica un subrayado simple
            field.Result.Font.Underline = win32com.client.constants.wdUnderlineSingle
            print(f"The hyperlink format has been changed: {field.Result.Text}")
    except Exception as e:
        print(f"Se produjo un error al procesar el campo: {e}")

# Recorrer todos los párrafos del documento en orden inverso
for i in range(doc.Paragraphs.Count, 0, -1):
    paragraph = doc.Paragraphs.Item(i)
    
    # Comprobar si el párrafo contiene una imagen
    if paragraph.Range.InlineShapes.Count > 0:
        # Si el párrafo anterior es un salto de párrafo, eliminarlo
        if i > 1:  # Asegurarse de que no es el primer párrafo
            prev_paragraph = doc.Paragraphs.Item(i - 1)
            if prev_paragraph.Range.Text.strip() == "":
               prev_paragraph.Range.Delete()





print("Reajustes de texto e imágenes saltos de párrafos")

# Constantes para identificar imágenes en Word
wdInlineShapePicture = 3  # Tipo de imagen para imágenes en línea "Picture" en InlineShape
POINTS_TO_CM = 0.0352778  # Conversión de puntos a centímetros

def limpiar_enumeraciones_vacias_antes_de_titulos(doc):
    print("Limpiando enumeraciones vacías antes de títulos...")

    i = 1
    while i <= doc.Paragraphs.Count:
        parrafo = doc.Paragraphs.Item(i)
        es_titulo = parrafo.OutlineLevel != win32.constants.wdOutlineLevelBodyText

        if es_titulo:
            j = i - 1
            while j > 0:
                parrafo_anterior = doc.Paragraphs.Item(j)
                es_enumeracion_anterior = parrafo_anterior.Range.ListFormat.ListType != win32.constants.wdListNoNumbering
                texto_anterior = parrafo_anterior.Range.Text.strip()

                if es_enumeracion_anterior and not texto_anterior:
                    parrafo_anterior.Range.Delete()
                    print(f"Enumeración vacía eliminada en la posición {j}.")
                    i -= 1
                else:
                    break
                j -= 1
        i += 1

def obtener_estilos_de_titulo(doc):
    print("Obteniendo títulos en el documento...")
    estilos_de_titulo = []
    for estilo in doc.Styles:
        if estilo.NameLocal.startswith("Heading") or "Heading" in estilo.NameLocal:
            estilos_de_titulo.append(estilo.NameLocal)
    i = 1
    while i <= doc.Paragraphs.Count:
        paragraph = doc.Paragraphs.Item(i)
        paragraph_text = paragraph.Range.Text.strip()
        style_name = paragraph.Style.NameLocal
        if style_name in estilos_de_titulo and paragraph_text:
            if i > 1:
                prev_paragraph = doc.Paragraphs.Item(i - 1)
                prev_text = prev_paragraph.Range.Text.strip()
                if prev_text != '':
                    paragraph.Range.InsertParagraphBefore()
                    new_paragraph = doc.Paragraphs.Item(i)
                    new_paragraph.Style = doc.Styles("Normal")
                    i += 1
            else:
                paragraph.Range.InsertParagraphBefore()
                new_paragraph = doc.Paragraphs.Item(1)
                new_paragraph.Style = doc.Styles("Normal")
                i += 1
            if i < doc.Paragraphs.Count:
                next_paragraph = doc.Paragraphs.Item(i + 1)
                next_text = next_paragraph.Range.Text.strip()
                if next_text != '':
                    paragraph.Range.InsertParagraphAfter()
                    new_paragraph = doc.Paragraphs.Item(i + 1)
                    new_paragraph.Style = doc.Styles("Normal")
                    i += 1
            else:
                paragraph.Range.InsertParagraphAfter()
                new_paragraph = doc.Paragraphs.Item(doc.Paragraphs.Count)
            i += 1
        else:
            i += 1

def mostrar_medidas_imagenes(doc):
    index = 1
    while index <= doc.InlineShapes.Count:
        inline_shape = doc.InlineShapes.Item(index)
        if inline_shape.Type == wdInlineShapePicture:
            width_cm = inline_shape.Width * POINTS_TO_CM
            height_cm = inline_shape.Height * POINTS_TO_CM

            paragraph = inline_shape.Range.Paragraphs(1)

            # Verificar si el párrafo está dentro de una tabla
            if paragraph.Range.Tables.Count > 0:
                # La imagen está dentro de una tabla
                # No podemos insertar párrafos antes o después dentro de una tabla
                print(f"La imagen en la posición {index} está dentro de una tabla. Se omite la inserción de párrafos.")
            else:
                # Obtener el índice del párrafo
                for i in range(1, doc.Paragraphs.Count + 1):
                    p = doc.Paragraphs.Item(i)
                    if p.Range.Start == paragraph.Range.Start:
                        paragraph_index = i
                        break

                if "|" in inline_shape.AlternativeText or "|" in inline_shape.Title:
                    paragraph.Range.InsertParagraphBefore()
                    paragraph.Range.InsertParagraphBefore()
                    paragraph_index += 2
                    paragraph = doc.Paragraphs.Item(paragraph_index)

                    new_paragraph_before = doc.Paragraphs.Item(paragraph_index - 1)
                    new_paragraph_before.Style = doc.Styles("Normal")
                    new_paragraph_before_2 = doc.Paragraphs.Item(paragraph_index - 2)
                    new_paragraph_before_2.Style = doc.Styles("Normal")

                    paragraph.Range.InsertParagraphAfter()
                    paragraph.Range.InsertParagraphAfter()

                    new_paragraph_after = doc.Paragraphs.Item(paragraph_index + 1)
                    new_paragraph_after.Style = doc.Styles("Normal")
                    new_paragraph_after_2 = doc.Paragraphs.Item(paragraph_index + 2)
                    new_paragraph_after_2.Style = doc.Styles("Normal")

                    print(f"Imagen con tubería detectada en la posición {index}. Se añadieron 2 saltos de párrafo antes y después.")
                else:
                    if paragraph_index > 1:
                        prev_paragraph = doc.Paragraphs.Item(paragraph_index - 1)
                        prev_text = prev_paragraph.Range.Text.strip()
                        if prev_text != '':
                            paragraph.Range.InsertParagraphBefore()
                            new_paragraph = doc.Paragraphs.Item(paragraph_index)
                            new_paragraph.Style = doc.Styles("Normal")
                            paragraph_index += 1
                            paragraph = doc.Paragraphs.Item(paragraph_index)
                    else:
                        paragraph.Range.InsertParagraphBefore()
                        new_paragraph = doc.Paragraphs.Item(1)
                        new_paragraph.Style = doc.Styles("Normal")
                        paragraph_index += 1
                        paragraph = doc.Paragraphs.Item(paragraph_index)

                    if paragraph_index < doc.Paragraphs.Count:
                        next_paragraph = doc.Paragraphs.Item(paragraph_index + 1)
                        next_text = next_paragraph.Range.Text.strip()
                        if next_text != '':
                            paragraph.Range.InsertParagraphAfter()
                            new_paragraph = doc.Paragraphs.Item(paragraph_index + 1)
                            new_paragraph.Style = doc.Styles("Normal")
                    else:
                        paragraph.Range.InsertParagraphAfter()
                        new_paragraph = doc.Paragraphs.Item(doc.Paragraphs.Count)
                        new_paragraph.Style = doc.Styles("Normal")
        index += 1

# Llamadas a las funciones para procesar el documento
limpiar_enumeraciones_vacias_antes_de_titulos(doc)
obtener_estilos_de_titulo(doc)
mostrar_medidas_imagenes(doc)



print("Paragraph Cleaning...")

# Limpieza de párrafos...
try:
    # Constantes para ComputeStatistics y GoTo
    wdStatisticPages = 2  # Valor para ComputeStatistics para páginas
    wdGoToPage = 1
    wdGoToAbsolute = 1

    # Comprobar si el documento tiene 6 páginas o más
    if doc.ComputeStatistics(wdStatisticPages) >= 6:
        sixth_page_range = word_app.Selection.GoTo(What=wdGoToPage, Which=wdGoToAbsolute, Count=6)
    else:
        sixth_page_range = None

    for i in range(doc.Paragraphs.Count, 0, -1):
        paragraph = doc.Paragraphs.Item(i)
        range_font = paragraph.Range.Font

        # Saltar párrafos antes de la sexta página, si existe
        if sixth_page_range and paragraph.Range.Start < sixth_page_range.Start:
            continue

        # Verificar si el párrafo es un título con la fuente "Graphik" (negrita, tamaño 16)
        if (
            range_font.Bold in [True, -1]  # Aceptar True o -1 como "negrita"
            and "Graphik" in range_font.Name  # Busca cualquier fuente que contenga "Graphik"
            and range_font.Size == 16
        ):
            # Verificar si el siguiente párrafo es una lista de viñetas
            next_paragraph = None
            if i < doc.Paragraphs.Count:
                next_paragraph = doc.Paragraphs.Item(i + 1)

            if next_paragraph and next_paragraph.Range.ListFormat.ListType in [1, 2]:  # Viñetas o numeración
                # Contar saltos de párrafo entre el título y la lista
                paragraph_gap_count = next_paragraph.Range.Start - paragraph.Range.End

                if paragraph_gap_count > 1:
                    # Si hay más de un salto de párrafo, reducir a uno
                    try:
                        doc.Range(paragraph.Range.End, next_paragraph.Range.Start).Delete()
                        paragraph.Range.InsertParagraphAfter()
                        print(f"Saltos de párrafo ajustados entre título en posición {i} y lista.")
                    except Exception as e:
                        print(f"Error ajustando saltos: {e}")

                elif paragraph_gap_count == 0:
                    # Si no hay salto de párrafo, insertar uno
                    try:
                        paragraph.Range.InsertParagraphAfter()
                        print(f"Salto de párrafo añadido entre título en posición {i} y lista.")
                    except Exception as e:
                        print(f"Error añadiendo salto: {e}")

        # Verificar si hay exactamente 4 saltos de párrafo consecutivos
        empty_paragraph_count = 0
        consecutive_paragraphs = []

        for j in range(i, doc.Paragraphs.Count + 1):
            consecutive_paragraph = doc.Paragraphs.Item(j)
            if consecutive_paragraph.Range.Text.strip() == "":
                empty_paragraph_count += 1
                consecutive_paragraphs.append(consecutive_paragraph)

                if empty_paragraph_count > 4:
                    break
            else:
                break

        if empty_paragraph_count == 4:
            try:
                # Eliminar dos párrafos para dejar solo dos
                doc.Range(consecutive_paragraphs[0].Range.Start, consecutive_paragraphs[1].Range.End).Delete()
                print(f"Saltos consecutivos reducidos de 4 a 2 en la posición {i}.")
            except Exception as e:
                print(f"Error eliminando saltos: {e}")

        # Verificar si hay exactamente 3 saltos de párrafo consecutivos
        empty_paragraph_count = 0
        consecutive_paragraphs = []

        for j in range(i, doc.Paragraphs.Count + 1):
            consecutive_paragraph = doc.Paragraphs.Item(j)
            if consecutive_paragraph.Range.Text.strip() == "":
                empty_paragraph_count += 1
                consecutive_paragraphs.append(consecutive_paragraph)

                if empty_paragraph_count > 3:
                    break
            else:
                break

        if empty_paragraph_count == 3:
            try:
                # Eliminar un párrafo para dejar solo dos
                doc.Range(consecutive_paragraphs[0].Range.Start, consecutive_paragraphs[0].Range.End).Delete()
                print(f"Saltos consecutivos reducidos de 3 a 2 en la posición {i}.")
            except Exception as e:
                print(f"Error eliminando salto: {e}")

except Exception as e:
    print(f"Error general: {e}")




print("Eliminando puntos negros vacios")

def eliminar_puntos_negros(doc):
    ## Itera sobre todos los párrafos del documento
    for para in doc.Paragraphs:
        # Obtén el texto del párrafo y elimina espacios adicionales
        texto_parrafo = para.Range.Text.strip()
        
        # Verifica si el párrafo es parte de una lista (viñetas o numeración)
        if para.Range.ListFormat.ListType != 0:  # No es "wdListNoNumbering"
            # Si el párrafo no tiene texto adicional aparte de la viñeta, lo elimina
            if not texto_parrafo or texto_parrafo == '•':
                try:
                    para.Range.Delete()  # Elimina el párrafo
                except Exception:
                    pass  # Ignora cualquier error al eliminar el párrafo

# Llama a la función (asegúrate de que 'doc' esté definido correctamente)
eliminar_puntos_negros(doc)



from datetime import datetime

print("Poniendo fecha actual al pie de página y la página del documento...")

def configurar_footer(doc):
    # Obtén el año actual
    year_actual = datetime.now().year

    try:
        # Asegúrate de que la Sección 2 existe
        if len(doc.Sections) >= 2:
            section = doc.Sections(2)  # Accede a la Sección 2

            # Habilitar pie de página diferente para la primera página
            section.PageSetup.DifferentFirstPageHeaderFooter = True

            # Desvincular el pie de página de la primera página de la sección anterior
            section.Footers(2).LinkToPrevious = False  # 2 = wdHeaderFooterFirstPage

            # Configura el pie de página para la primera página (First Page Footer)
            first_page_footer = section.Footers(2)  # 2 = wdHeaderFooterFirstPage
            first_page_footer.Range.Text = ""  # Limpia cualquier contenido previo en el pie de página

            # Configurar el Copyright centrado
            rango_copyright = first_page_footer.Range
            rango_copyright.ParagraphFormat.Alignment = 1  # Centrado
            rango_copyright.Text = f"Copyright © {year_actual} Accenture. All rights reserved."
            rango_copyright.Font.Bold = False
            rango_copyright.Font.Name = "Arial"
            rango_copyright.Font.Size = 10

            # Configurar el número de página alineado a la derecha
            rango_numero_pagina = first_page_footer.Range.Duplicate
            rango_numero_pagina.Collapse(0)  # Colapsa al final del contenido actual
            rango_numero_pagina.Fields.Add(rango_numero_pagina, 33)  # 33 = wdFieldPage
            rango_numero_pagina.InsertBefore("\t\t\t\t\t\t")  # Inserta 5 tabulaciones
            rango_numero_pagina.ParagraphFormat.Alignment = 2  # Derecha

            # Configurar el inicio de la numeración desde esta sección
            section.Headers(1).PageNumbers.RestartNumberingAtSection = True
            section.Headers(1).PageNumbers.StartingNumber = 1

        else:
            print("La Sección 2 no existe en el documento.")

        # Actualiza todos los campos dinámicos del documento
        doc.Fields.Update()
        print("Campos del documento actualizados correctamente.")

    except Exception as e:
        print(f"Error al configurar el pie de página: {e}")

# Llama a la función (asegúrate de que 'doc' esté definido correctamente)
configurar_footer(doc)


try:
    # Verificar si existe al menos una tabla de contenido
    if doc.TablesOfContents.Count > 0:
        # Acceder a la tabla de contenido
        table_of_contents = doc.TablesOfContents(1)
        try:
            # Intentar actualizar la tabla de contenido
            table_of_contents.Update()
            print("Tabla de contenido actualizada correctamente.")
        except Exception as e:
            print(f"Se produjo un error al intentar actualizar la tabla de contenido: {e}")
    else:
        print("No se encontró ninguna tabla de contenido en el documento.")
except Exception as e:
    print(f"Se produjo un error al intentar acceder a la tabla de contenido: {e}")



try:
    process_document()
    doc.Save()
    print("Document processed and saved successfully.")
except Exception as e:
    print(f"An error occurred while processing the document: {e}")

try:
    doc.Close()
    word_app.Quit()
    print("Word application closed.")
except Exception as e:
    print(f"An error occurred while closing Word: {e}")

print("Finished!")

    """)

# Activar el entorno virtual "shell" solo si no está activado
if 'VIRTUAL_ENV' not in os.environ:
    activate_script = os.path.join("shell", "Scripts", "activate")
    if sys.platform == "linux":
        activate_script = os.path.join("shell", "bin", "activate")

    # En Windows, necesitas ejecutar el script de activación en una shell
    if sys.platform == "win32":
        command = f"{activate_script} && {venv_python} second_script.py"
        subprocess.run(["cmd", "/k", command])
    else:  # En Unix, puedes usar 'source'
        command = f"source {activate_script} && {venv_python} second_script.py"
        subprocess.run(["bash", "-c", command])
